#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
tutorial_docx_builder.py
Generador de tutoriales IT en formato DOCX, estilo O'Reilly / Perplexity Computer.
Paleta de colores y estructura derivadas del análisis de PDFs de referencia.

Uso:
    from tutorial_docx_builder import generate_tutorial_docx
    generate_tutorial_docx(content_dict, "/ruta/output.docx")
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta de Colores ─────────────────────────────────────────────────────────
TEAL_COVER    = '0F3D3D'   # Fondo portada (verde-azul oscuro)
TEAL_ACCENT   = '1D9CB8'   # Encabezados de fase, prefijo Paso, sub-secciones
TEAL_TBL_HDR  = '1A5A6A'   # Fila de encabezado de tabla
CODE_BG       = '1A2233'   # Fondo de bloque de código
CODE_FG       = 'E8EAF0'   # Texto de código
CODE_COMMENT  = '7EC8A0'   # Comentarios en código
ALT_ROW       = 'EEF4F6'   # Fila alternante de tabla

CALLOUTS = {
    'IMPORTANTE':  ('D97706', 'FFFBEB'),
    'NOTA':        ('2563EB', 'EFF6FF'),
    'CONSEJO':     ('16A34A', 'F0FDF4'),
    'ADVERTENCIA': ('9333EA', 'FAF5FF'),
    'ATENCIÓN':    ('DC2626', 'FEF2F2'),
    'ATENCION':    ('DC2626', 'FEF2F2'),
}


# ── Utilidades XML ────────────────────────────────────────────────────────────

def _hex_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    # Eliminar shd previo si existe
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _cell_left_border(cell, color_hex, sz='24'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        if side == 'left':
            b.set(qn('w:val'), 'thick')
            b.set(qn('w:sz'), sz)
            b.set(qn('w:color'), color_hex)
        else:
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:color'), 'auto')
        b.set(qn('w:space'), '0')
        borders.append(b)
    tc_pr.append(borders)


def _table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _cell_padding(cell, top=120, left=200, bottom=120, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tc_pr.append(tcMar)


def _para_bottom_border(paragraph, color='CCCCCC', sz=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _page_number_field(run):
    for tag, text in [('begin', None), ('instrText', 'PAGE'), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.text = text
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)


# ── Funciones de construcción del documento ───────────────────────────────────

def create_document():
    """Crea un documento A4 con márgenes y estilos base."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(8)
    return doc


def add_cover_page(doc, title, subtitle, version_line, specs, date=''):
    """
    Portada de página completa con fondo verde-azul oscuro (estilo Perplexity).
    specs: dict ordenado {etiqueta: valor}
    """
    tbl = doc.add_table(rows=1, cols=1)
    _table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    _cell_bg(cell, TEAL_COVER)

    # Altura mínima: casi toda la página
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(Cm(22.5).twips)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)
    _cell_padding(cell, top=600, left=480, bottom=400, right=480)

    # Título (blanco, 38pt, sin negrita)
    p_t = cell.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_t.paragraph_format.space_before = Pt(48)
    p_t.paragraph_format.space_after  = Pt(10)
    r = p_t.add_run(title)
    r.font.name = 'Calibri Light'
    r.font.size = Pt(38)
    r.font.bold = False
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Subtítulo (color acento teal)
    p_s = cell.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_s.paragraph_format.space_before = Pt(0)
    p_s.paragraph_format.space_after  = Pt(6)
    r_s = p_s.add_run(subtitle)
    r_s.font.name = 'Calibri'
    r_s.font.size = Pt(17)
    r_s.font.color.rgb = _hex_rgb(TEAL_ACCENT)

    # Línea de versión (gris claro)
    p_v = cell.add_paragraph()
    p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_v.paragraph_format.space_before = Pt(0)
    p_v.paragraph_format.space_after  = Pt(28)
    r_v = p_v.add_run(version_line)
    r_v.font.name = 'Calibri'
    r_v.font.size = Pt(11)
    r_v.font.color.rgb = RGBColor(0xB0, 0xBE, 0xC5)

    # Separador visual (línea verde-lima)
    p_sep = cell.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(0)
    p_sep.paragraph_format.space_after  = Pt(18)
    _para_bottom_border(p_sep, '76C442', sz=12)

    # Bloque de especificaciones (monoespacio verde)
    if specs:
        p_sh = cell.add_paragraph()
        p_sh.paragraph_format.space_before = Pt(8)
        p_sh.paragraph_format.space_after  = Pt(6)
        r_sh = p_sh.add_run('Especificaciones del Sistema')
        r_sh.font.name = 'Calibri'
        r_sh.font.size = Pt(9)
        r_sh.font.bold = True
        r_sh.font.color.rgb = RGBColor(0x80, 0x9A, 0x80)

        for lbl, val in specs.items():
            p_sp = cell.add_paragraph()
            p_sp.paragraph_format.space_before = Pt(0)
            p_sp.paragraph_format.space_after  = Pt(2)
            r_lbl = p_sp.add_run(f'{lbl:<20}: ')
            r_lbl.font.name = 'Courier New'
            r_lbl.font.size = Pt(9)
            r_lbl.font.color.rgb = RGBColor(0x7E, 0xC8, 0x9A)
            r_val = p_sp.add_run(val)
            r_val.font.name = 'Courier New'
            r_val.font.size = Pt(9)
            r_val.font.color.rgb = RGBColor(0xCC, 0xE8, 0xD8)

    # Pie de portada
    p_d = cell.add_paragraph()
    p_d.paragraph_format.space_before = Pt(32)
    r_d = p_d.add_run(f'Generado con Claude Code  ·  {date}')
    r_d.font.name = 'Calibri'
    r_d.font.size = Pt(9)
    r_d.font.color.rgb = RGBColor(0x60, 0x7A, 0x7A)

    doc.add_page_break()


def add_toc_page(doc, chapters, appendices=None):
    """Página de tabla de contenidos estilo Perplexity."""
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(28)
    r = p_title.add_run('Tabla de Contenidos')
    r.font.name = 'Calibri Light'
    r.font.size = Pt(26)
    r.font.color.rgb = _hex_rgb(TEAL_ACCENT)

    for ch in chapters:
        lbl = ch.get('toc_label', f'Fase {ch["number"]}')
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        r_n = p.add_run(f'{lbl}:  ')
        r_n.font.name = 'Calibri'
        r_n.font.size = Pt(11)
        r_n.font.bold = True
        r_n.font.color.rgb = _hex_rgb(TEAL_ACCENT)
        r_ttl = p.add_run(ch['title'])
        r_ttl.font.name = 'Calibri'
        r_ttl.font.size = Pt(11)

        if ch.get('description'):
            p_d = doc.add_paragraph(ch['description'])
            p_d.paragraph_format.left_indent = Cm(1.5)
            p_d.paragraph_format.space_before = Pt(1)
            p_d.paragraph_format.space_after  = Pt(6)
            for r_d in p_d.runs:
                r_d.font.size = Pt(9)
                r_d.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        _hr(doc, color='DDDDDD', sz=4)

    if appendices:
        for app in appendices:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            r_a = p.add_run(f'Apéndice {app["letter"]}:  ')
            r_a.font.name = 'Calibri'
            r_a.font.size = Pt(11)
            r_a.font.bold = True
            r_a.font.color.rgb = _hex_rgb(TEAL_ACCENT)
            r_at = p.add_run(app['title'])
            r_at.font.name = 'Calibri'
            r_at.font.size = Pt(11)

    doc.add_page_break()


def _hr(doc, color='AAAAAA', sz=8):
    """Línea horizontal delgada entre secciones."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    _para_bottom_border(p, color, sz)
    return p


def add_chapter_heading(doc, number, title, label='Fase'):
    """Encabezado de capítulo/fase grande con regla inferior."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'{label} {number}: {title}')
    r.font.name = 'Calibri Light'
    r.font.size = Pt(28)
    r.font.bold = False
    r.font.color.rgb = RGBColor(0x18, 0x18, 0x18)
    _hr(doc, 'AAAAAA', 10)


def add_objective_box(doc, text, label='Objetivo de esta fase'):
    """Caja con borde izquierdo teal y fondo azul muy claro."""
    tbl = doc.add_table(rows=1, cols=1)
    _table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    _cell_bg(cell, 'E4F3F7')
    _cell_left_border(cell, TEAL_ACCENT, sz='18')
    _cell_padding(cell, top=120, left=220, bottom=120, right=160)

    p_lbl = cell.paragraphs[0]
    p_lbl.paragraph_format.space_before = Pt(0)
    p_lbl.paragraph_format.space_after  = Pt(4)
    r_lbl = p_lbl.add_run(label)
    r_lbl.font.name = 'Calibri'
    r_lbl.font.size = Pt(9)
    r_lbl.font.bold = True
    r_lbl.font.color.rgb = _hex_rgb(TEAL_ACCENT)

    p_txt = cell.add_paragraph(text)
    p_txt.paragraph_format.space_before = Pt(0)
    p_txt.paragraph_format.space_after  = Pt(0)
    p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p_txt.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

    _spacer(doc)


def add_subsection(doc, title):
    """Sub-sección con título en color teal y regla inferior ligera."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(title)
    r.font.name = 'Calibri'
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = _hex_rgb(TEAL_ACCENT)
    _hr(doc, 'CCEAF0', 4)


def add_step_heading(doc, number, title):
    """Paso N: Título — prefijo teal, texto negro."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r_pre = p.add_run(f'Paso {number}: ')
    r_pre.font.name = 'Calibri'
    r_pre.font.size = Pt(13)
    r_pre.font.bold = True
    r_pre.font.color.rgb = _hex_rgb(TEAL_ACCENT)
    r_ttl = p.add_run(title)
    r_ttl.font.name = 'Calibri'
    r_ttl.font.size = Pt(13)
    r_ttl.font.bold = False
    r_ttl.font.color.rgb = RGBColor(0x18, 0x18, 0x18)


def add_body_text(doc, text):
    """
    Párrafo de cuerpo justificado.
    Soporta fragmentos inline con ``código`` delimitados por backticks dobles.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)

    # Dividir por código inline (delimitado con ` o ``)
    import re
    parts = re.split(r'`+([^`]+)`+', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if i % 2 == 1:  # Es fragmento de código inline
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x6A, 0x8A)


def add_code_block(doc, code):
    """Bloque de código oscuro con sintaxis monoespacio."""
    tbl = doc.add_table(rows=1, cols=1)
    _table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    _cell_bg(cell, CODE_BG)
    _cell_padding(cell, top=140, left=220, bottom=140, right=160)

    lines = code.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        stripped = line.lstrip()
        if stripped.startswith('#') or stripped.startswith('//') or stripped.startswith('--'):
            run.font.color.rgb = _hex_rgb(CODE_COMMENT)
        else:
            run.font.color.rgb = _hex_rgb(CODE_FG)

    _spacer(doc)


def add_expected_output(doc, output, label='Salida esperada:'):
    """Bloque de salida esperada — ligeramente más oscuro que el código."""
    p_lbl = doc.add_paragraph(label)
    p_lbl.paragraph_format.space_before = Pt(6)
    p_lbl.paragraph_format.space_after  = Pt(3)
    p_lbl.runs[0].font.name = 'Calibri'
    p_lbl.runs[0].font.size = Pt(10)
    p_lbl.runs[0].font.bold = False

    tbl = doc.add_table(rows=1, cols=1)
    _table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    _cell_bg(cell, '0D1520')
    _cell_padding(cell, top=120, left=200, bottom=120, right=140)

    for i, line in enumerate(output.split('\n')):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xA0, 0xE0, 0xC0)

    _spacer(doc)


def add_callout(doc, callout_type, content):
    """
    Caja de aviso con borde izquierdo de color.
    Tipos: IMPORTANTE, NOTA, CONSEJO, ADVERTENCIA, ATENCIÓN
    """
    key = callout_type.upper().replace('Ó', 'O').replace('É', 'E')
    border_color, bg_color = CALLOUTS.get(key, CALLOUTS.get(callout_type.upper(), ('888888', 'F8F8F8')))

    tbl = doc.add_table(rows=1, cols=1)
    _table_no_borders(tbl)
    cell = tbl.cell(0, 0)
    _cell_bg(cell, bg_color)
    _cell_left_border(cell, border_color, sz='24')
    _cell_padding(cell, top=120, left=220, bottom=120, right=160)

    p_lbl = cell.paragraphs[0]
    p_lbl.paragraph_format.space_before = Pt(0)
    p_lbl.paragraph_format.space_after  = Pt(5)
    r_lbl = p_lbl.add_run(callout_type.upper())
    r_lbl.font.name = 'Calibri'
    r_lbl.font.size = Pt(9)
    r_lbl.font.bold = True
    r_lbl.font.color.rgb = _hex_rgb(border_color)

    p_txt = cell.add_paragraph(content)
    p_txt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_txt.paragraph_format.space_before = Pt(0)
    p_txt.paragraph_format.space_after  = Pt(0)
    for r in p_txt.runs:
        r.font.name = 'Calibri'
        r.font.size = Pt(10)

    _spacer(doc)


def add_table(doc, headers, rows, caption=None):
    """Tabla con encabezado teal oscuro y filas alternas."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fila de encabezado
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _cell_bg(cell, TEAL_TBL_HDR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Filas de datos
    for ri, row_data in enumerate(rows):
        bg = ALT_ROW if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            _cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.name = 'Calibri'
            r.font.size = Pt(10)

    if caption:
        p_cap = doc.add_paragraph(caption)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after  = Pt(10)
        p_cap.runs[0].font.name = 'Calibri'
        p_cap.runs[0].font.size = Pt(9)
        p_cap.runs[0].font.italic = True
        p_cap.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        _spacer(doc)


def add_bullet_list(doc, items, numbered=False):
    """Lista de viñetas o numerada."""
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        prefix = f'{i + 1}. ' if numbered else '• '
        run = p.add_run(prefix + item)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    _spacer(doc, Pt(4))


def add_appendix_heading(doc, letter, title):
    """Encabezado de apéndice."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    r_lbl = p.add_run(f'APÉNDICE {letter}: ')
    r_lbl.font.name = 'Calibri'
    r_lbl.font.size = Pt(22)
    r_lbl.font.bold = True
    r_lbl.font.color.rgb = _hex_rgb(TEAL_ACCENT)
    r_ttl = p.add_run(title)
    r_ttl.font.name = 'Calibri Light'
    r_ttl.font.size = Pt(22)
    r_ttl.font.color.rgb = RGBColor(0x18, 0x18, 0x18)
    _hr(doc, 'AAAAAA', 10)


def _spacer(doc, height=None):
    """Párrafo vacío de separación."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = height if height else Pt(8)


def setup_header_footer(doc, doc_title, product, version, date):
    """Configura encabezado y pie en todas las páginas excepto la portada."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True

    # ── Encabezado ────────────────────────────────────────────────────
    hdr = sec.header
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

    # Configurar tabulación derecha
    from docx.oxml import OxmlElement as OX
    pPr = p._p.get_or_add_pPr()
    tabs = OX('w:tabs')
    tab = OX('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(Cm(16.0).twips)))
    tabs.append(tab)
    pPr.append(tabs)

    r1 = p.add_run(f'{doc_title}  |  {product}')
    r1.font.name = 'Calibri'
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run('\t').font.size = Pt(8)
    r2 = p.add_run(f'{version}  |  {date}')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _para_bottom_border(p, 'CCCCCC', 4)

    # ── Pie de página ─────────────────────────────────────────────────
    ftr = sec.footer
    p_f = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    p_f.clear()
    p_f.paragraph_format.space_before = Pt(4)
    p_f.paragraph_format.space_after  = Pt(0)

    pPr2 = p_f._p.get_or_add_pPr()
    tabs2 = OX('w:tabs')
    tab2 = OX('w:tab')
    tab2.set(qn('w:val'), 'right')
    tab2.set(qn('w:pos'), str(int(Cm(16.0).twips)))
    tabs2.append(tab2)
    pPr2.append(tabs2)

    r_f = p_f.add_run(f'Generado con Claude Code  |  {date}')
    r_f.font.name = 'Calibri'
    r_f.font.size = Pt(8)
    r_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p_f.add_run('\t').font.size = Pt(8)

    r_pg = p_f.add_run('Página ')
    r_pg.font.name = 'Calibri'
    r_pg.font.size = Pt(8)
    r_pg.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    r_num = p_f.add_run()
    r_num.font.name = 'Calibri'
    r_num.font.size = Pt(8)
    r_num.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    _page_number_field(r_num)


# ── Motor de renderizado de secciones ─────────────────────────────────────────

def _render_section(doc, section):
    """Renderiza un elemento de contenido según su tipo."""
    t = section.get('type', 'body')

    if t == 'body':
        add_body_text(doc, section['content'])
    elif t == 'subsection':
        add_subsection(doc, section['title'])
    elif t == 'step':
        add_step_heading(doc, section['number'], section['title'])
        if section.get('body'):
            add_body_text(doc, section['body'])
    elif t == 'code':
        add_code_block(doc, section['content'])
    elif t == 'output':
        add_expected_output(doc, section['content'], section.get('label', 'Salida esperada:'))
    elif t == 'callout':
        add_callout(doc, section['callout_type'], section['content'])
    elif t == 'table':
        add_table(doc, section['headers'], section['rows'], section.get('caption'))
    elif t == 'bullets':
        add_bullet_list(doc, section['items'], section.get('numbered', False))


# ── Punto de entrada principal ────────────────────────────────────────────────

def generate_tutorial_docx(content, output_path):
    """
    Genera un archivo .docx completo a partir del diccionario de contenido.

    Estructura esperada de `content`:
    {
        "title":    "Título del Tutorial",
        "subtitle": "Producto / Contexto",
        "version":  "JetPack 6.2.2 | Ubuntu 22.04 | CUDA 12.6",
        "date":     "Abril 2026",
        "specs":    {"Hardware": "...", "OS": "..."},  # dict ordenado
        "chapters": [
            {
                "number":      1,
                "title":       "Nombre de la fase",
                "description": "Breve descripción para el TOC",
                "label":       "Fase",   # o "Capítulo"
                "toc_label":   "Fase 1", # etiqueta en el TOC
                "objective":   "Texto del objetivo (opcional)",
                "sections": [
                    {"type": "body",       "content": "..."},
                    {"type": "subsection", "title":   "..."},
                    {"type": "step",       "number": "1", "title": "...", "body": "..."},
                    {"type": "code",       "content": "..."},
                    {"type": "output",     "content": "...", "label": "Salida esperada:"},
                    {"type": "callout",    "callout_type": "IMPORTANTE", "content": "..."},
                    {"type": "table",      "headers": [...], "rows": [[...]], "caption": "..."},
                    {"type": "bullets",    "items": [...], "numbered": False},
                ]
            }
        ],
        "appendices": [
            {
                "letter": "A",
                "title":  "Referencia Rápida de Comandos",
                "sections": [...]
            }
        ]
    }
    """
    doc = create_document()

    # 1. Portada
    add_cover_page(
        doc,
        content['title'],
        content.get('subtitle', ''),
        content.get('version', ''),
        content.get('specs', {}),
        content.get('date', '')
    )

    # 2. Tabla de contenidos
    add_toc_page(doc, content.get('chapters', []), content.get('appendices'))

    # 3. Encabezado y pie (excluye portada gracias a different_first_page)
    setup_header_footer(
        doc,
        content['title'],
        content.get('subtitle', ''),
        content.get('version', ''),
        content.get('date', '')
    )

    # 4. Capítulos / Fases
    for chapter in content.get('chapters', []):
        add_chapter_heading(
            doc,
            chapter['number'],
            chapter['title'],
            chapter.get('label', 'Fase')
        )
        if chapter.get('objective'):
            add_objective_box(doc, chapter['objective'])
        for section in chapter.get('sections', []):
            _render_section(doc, section)
        doc.add_page_break()

    # 5. Apéndices
    for app in content.get('appendices', []):
        add_appendix_heading(doc, app['letter'], app['title'])
        for section in app.get('sections', []):
            _render_section(doc, section)
        doc.add_page_break()

    doc.save(output_path)
    print(f'✓ Tutorial generado: {output_path}')
    return output_path
