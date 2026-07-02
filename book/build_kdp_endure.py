#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_kdp_endure.py — Convierte capítulos Markdown a DOCX con estilos Amazon Endure KDP (7×10 pulgadas).

Uso individual:
    python book/build_kdp_endure.py book/parte-a/capitulo-01-inicio-rapido/capitulo-01-inicio-rapido.md

Uso batch (genera 42 DOCX individuales para revisión):
    python book/build_kdp_endure.py --batch

Uso compilación final (tras revisión del autor):
    python book/build_kdp_endure.py --compile [salida.docx]

Requiere: pip install python-docx matplotlib pillow
"""

import sys
import re
import os
import io
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch

# ── Rutas ────────────────────────────────────────────────────────────────────
BOOK_DIR       = Path(__file__).parent
TEMPLATE_PATH  = Path('/root/.claude/uploads/a54997a3-f611-5b20-8996-cfdb2ca116be/9cf1d737-7_x_10_in_Spanish.docx')
INFOGRAFIAS_DIR = BOOK_DIR / 'infografias'
DOCX_OUT_DIR    = BOOK_DIR / 'docx-por-capitulo'

# ── Paleta NVIDIA ─────────────────────────────────────────────────────────────
NVIDIA_DARK    = '#0F3D3D'
NVIDIA_ACCENT  = '#1D9CB8'
NVIDIA_TABLE   = '#1A5A6A'
CODE_BG        = '#EAEAEA'   # obs#7: changed from F2F2F2
CODE_BORDER    = '#1A5A6A'
TABLE_BG       = 'EAEAEA'   # obs#6: uniform data row bg
TABLE_BORDER   = 'FFFFFF'   # obs#6: white borders
TABLE_BORDER_SZ = '12'      # obs#6: 1.5pt in OOXML eighths
CODE_FG        = '#1A1A1A'
CODE_COMMENT   = '#888888'
CODE_OUTPUT    = '#555555'
CALLOUT_WARN   = '#C0392B'
CALLOUT_IMP    = '#E67E22'
CALLOUT_TIP    = '#1A5A6A'
CALLOUT_NOTE   = '#888888'

# ── Nombres de estilos KDP Endure ─────────────────────────────────────────────
STYLE_CHAPTER_TITLE   = 'Endure - Chapter Title'
STYLE_BODY            = 'Endure - Chapter Body Text'
STYLE_FIRST_PARA      = 'Endure - First Paragraph Body Text'
STYLE_SUBHEAD         = 'Endure - Subhead'
STYLE_FRONT_MATTER    = 'Endure - Front Matter Body Text'
STYLE_COPYRIGHT       = 'Endure - Copyright Page'

# ── Orden canónico del libro completo ────────────────────────────────────────
BOOK_CHAPTERS = [
    'prologo/prologo.md',
    'introduccion/introduccion.md',
    'parte-a/00-parte-a/parte-a.md',
    'parte-a/capitulo-01-inicio-rapido/capitulo-01-inicio-rapido.md',
    'parte-a/capitulo-02-primer-arranque/capitulo-02-primer-arranque.md',
    'parte-a/capitulo-03-configuracion-base/capitulo-03-configuracion-base.md',
    'parte-a/capitulo-04-memoria-almacenamiento/capitulo-04-memoria-almacenamiento.md',
    'parte-a/capitulo-05-rendimiento/capitulo-05-rendimiento.md',
    'parte-a/capitulo-06-shell-entorno/capitulo-06-shell-entorno.md',
    'parte-a/capitulo-07-red/capitulo-07-red.md',
    'parte-a/capitulo-08-acceso-remoto/capitulo-08-acceso-remoto.md',
    'parte-a/capitulo-09-docker/capitulo-09-docker.md',
    'parte-a/capitulo-10-motores-inferencia/capitulo-10-motores-inferencia.md',
    'parte-a/capitulo-11-stack-agentico/capitulo-11-stack-agentico.md',
    'parte-a/capitulo-11a-openclaw/capitulo-11a-openclaw.md',
    'parte-a/capitulo-11b-nemoclaw/capitulo-11b-nemoclaw.md',
    'parte-a/capitulo-11c-open-webui/capitulo-11c-open-webui.md',
    'parte-a/capitulo-11d-tool-calling/capitulo-11d-tool-calling.md',
    'parte-a/capitulo-12-computer-vision/capitulo-12-computer-vision.md',
    'parte-a/capitulo-13-tts-stt/capitulo-13-tts-stt.md',
    'parte-a/capitulo-14-imagen-video/capitulo-14-imagen-video.md',
    'parte-a/capitulo-15-n8n/capitulo-15-n8n.md',
    'parte-a/capitulo-16-benchmarking/capitulo-16-benchmarking.md',
    'parte-a/capitulo-17-produccion/capitulo-17-produccion.md',
    'parte-a/capitulo-18-troubleshooting/capitulo-18-troubleshooting.md',
    'parte-b/00-parte-b/parte-b.md',
    'parte-b/capitulo-19-python-vscode/capitulo-19-python-vscode.md',
    'parte-b/capitulo-20-jetson-containers/capitulo-20-jetson-containers.md',
    'parte-b/capitulo-21-pdf-podcast/capitulo-21-pdf-podcast.md',
    'parte-b/capitulo-22-transcripcion-reuniones/capitulo-22-transcripcion-reuniones.md',
    'parte-b/capitulo-23-agencia-turismo/capitulo-23-agencia-turismo.md',
    'parte-b/capitulo-24-embudo-ventas/capitulo-24-embudo-ventas.md',
    'parte-b/capitulo-25-linkedin/capitulo-25-linkedin.md',
    'parte-b/capitulo-26-asistente-voz/capitulo-26-asistente-voz.md',
    'parte-b/capitulo-27-rag/capitulo-27-rag.md',
    'parte-b/capitulo-28-microservicios-saas/capitulo-28-microservicios-saas.md',
    'parte-c/00-parte-c/parte-c.md',
    'parte-c/capstone-01-agencia-ia/capstone-01-agencia-ia.md',
    'parte-c/capstone-02-automatizacion-video/capstone-02-automatizacion-video.md',
    'parte-c/conclusiones/conclusiones.md',
    'glosario/glosario.md',
    'parte-a/appendix/appendix.md',
]


# ═══════════════════════════════════════════════════════════════════════════════
# UTILIDADES DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def hex_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.lstrip('#'))
    tc_pr.append(shd)


def table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def set_left_border(cell, color_hex, sz='24', val='thick'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), val)
    left.set(qn('w:sz'), sz)
    left.set(qn('w:color'), color_hex.lstrip('#'))
    borders.append(left)
    tc_pr.append(borders)


def set_full_border(cell, color_hex, sz='4', left_sz='12'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color_hex.lstrip('#'))
        borders.append(el)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), left_sz)
    left.set(qn('w:color'), color_hex.lstrip('#'))
    borders.append(left)
    tc_pr.append(borders)


def set_cell_padding(cell, top=80, left=160, bottom=80, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tc_mar.append(m)
    tc_pr.append(tc_mar)


def set_all_cell_borders(cell, color_hex, sz):
    """Aplica bordes del mismo color a todos los lados de una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:tcBorders')):
        tc_pr.remove(old)
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color_hex.lstrip('#'))
        borders.append(el)
    tc_pr.append(borders)


def set_row_height(row, cm_value):
    """Altura mínima de fila en cm (1 cm ≈ 567 twips)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(cm_value * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def set_repeat_header(row):
    """Marca la fila para repetirse como encabezado en cada página."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)


def set_cell_valign(cell, val='center'):
    """Alineación vertical de celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:vAlign')):
        tc_pr.remove(old)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tc_pr.append(vAlign)


def set_outline_level(para, level):
    """Establece el nivel de esquema para TOC automático de Word (1=H1, 2=H2, 3=H3)."""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:outlineLvl')):
        pPr.remove(old)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level - 1))  # 0-indexed
    pPr.append(outlineLvl)


def keep_together(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kt = OxmlElement('w:keepLines')
    pPr.append(kt)
    kt2 = OxmlElement('w:keepNext')
    pPr.append(kt2)


def style_exists(doc, style_name):
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False


# ═══════════════════════════════════════════════════════════════════════════════
# CREACIÓN DEL DOCUMENTO BASE
# ═══════════════════════════════════════════════════════════════════════════════

def create_chapter_doc():
    """Abre la plantilla KDP y limpia el contenido, preservando todos los estilos Endure."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f'Plantilla KDP no encontrada: {TEMPLATE_PATH}')
    doc = Document(str(TEMPLATE_PATH))
    # Eliminar todos los párrafos existentes del cuerpo
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            body.remove(child)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# RENDERIZADO CON ESTILOS ENDURE
# ═══════════════════════════════════════════════════════════════════════════════

def get_style(doc, preferred, fallback='Normal'):
    return preferred if style_exists(doc, preferred) else fallback


def add_chapter_title(doc, text):
    style = get_style(doc, STYLE_CHAPTER_TITLE, 'Heading 1')
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    # Detectar si hay número de capítulo y separarlo
    if ' — ' in text:
        parts = text.split(' — ', 1)
        r1 = p.add_run(parts[0] + ' — ')
        r1.bold = False
        r2 = p.add_run(parts[1])
        r2.bold = False
    else:
        p.add_run(text)
    set_outline_level(p, 1)  # obs#11: TOC support
    return p


def add_part_title(doc, text, subtitle=''):
    style = get_style(doc, 'Heading 1')
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    if subtitle:
        ps = doc.add_paragraph(style=get_style(doc, 'Heading 2'))
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after = Pt(18)
        ps.add_run(subtitle)
    return p


def add_subhead(doc, text, level=2):
    style = get_style(doc, STYLE_SUBHEAD, 'Heading 2')
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    if level == 2:
        r.font.size = Pt(12)  # obs#5
    elif level == 3:
        r.font.size = Pt(11)  # obs#5: was 9pt
    set_outline_level(p, level)  # obs#11: TOC support
    return p


def add_subsubhead(doc, text):
    p = doc.add_paragraph(style=get_style(doc, STYLE_BODY, 'Normal'))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    return p


def add_first_paragraph(doc, text, is_front_matter=False):
    style = STYLE_FRONT_MATTER if is_front_matter else STYLE_FIRST_PARA
    style = get_style(doc, style, 'Normal')
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(5)   # obs#1
    p.paragraph_format.space_after  = Pt(5)   # obs#1
    _add_inline_formatted_runs(p, text, font_size=Pt(10))  # obs#5
    return p


def add_body_paragraph(doc, text, is_front_matter=False):
    style = STYLE_FRONT_MATTER if is_front_matter else STYLE_BODY
    style = get_style(doc, style, 'Normal')
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(5)   # obs#1
    p.paragraph_format.space_after  = Pt(5)   # obs#1
    _add_inline_formatted_runs(p, text, font_size=Pt(10))  # obs#5
    return p


def _add_inline_formatted_runs(paragraph, text, font_size=None):
    """Procesa markdown inline: **bold**, *italic*, `code`, [link](url)."""
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            r = paragraph.add_run(text[last:m.start()])
            if font_size:
                r.font.size = font_size
        full = m.group(0)
        if full.startswith('**'):
            r = paragraph.add_run(m.group(2))
            r.bold = True
            if font_size:
                r.font.size = font_size
        elif full.startswith('*'):
            r = paragraph.add_run(m.group(3))
            r.italic = True
            if font_size:
                r.font.size = font_size
        elif full.startswith('`'):
            r = paragraph.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(8)   # obs#5: inline code 8pt
        last = m.end()
    if last < len(text):
        r = paragraph.add_run(text[last:])
        if font_size:
            r.font.size = font_size


def split_code_at_separators(code_text):
    """Divide código en sub-bloques en líneas '# ── SectionName ──────'."""
    SEP = re.compile(r'^#\s*[─—\-]{3,}')
    sections, current = [], []
    for line in code_text.split('\n'):
        if SEP.match(line.strip()):
            if any(ln.strip() for ln in current):
                sections.append('\n'.join(current))
            current = [line]
        else:
            current.append(line)
    if any(ln.strip() for ln in current):
        sections.append('\n'.join(current))
    return sections if len(sections) > 1 else [code_text]


def _render_code_table(doc, code, lang='bash'):
    """Renderiza un bloque de código como tabla 1×1."""
    tbl = doc.add_table(rows=1, cols=1)
    table_full_width(tbl)
    cell = tbl.cell(0, 0)
    cell_bg(cell, CODE_BG)
    set_left_border(cell, CODE_BORDER, sz='24', val='thick')
    set_cell_padding(cell, top=100, left=180, bottom=100, right=140)

    lines = code.split('\n')
    while lines and not lines[-1].strip():
        lines.pop()

    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        keep_together(p)

        stripped = line.strip()
        # obs#7: comments in ALL languages, not just bash/js
        is_comment = stripped.startswith('#') or stripped.startswith('//')
        is_output = (
            lang in ('bash', 'sh', 'shell', '')
            and not is_comment
            and not stripped.startswith('$')
            and bool(stripped)
            and not re.match(r'^[A-Z_]+=|^[a-z_]+=', stripped)
            and not any(stripped.startswith(cmd) for cmd in (
                'sudo', 'git', 'pip', 'pip3', 'docker', 'python', 'python3',
                'curl', 'apt', 'apt-get', 'echo', 'cat', 'export', 'cp',
                'mv', 'rm', 'mkdir', 'ls', 'cd', 'nano', 'vim', 'source',
                'chmod', 'chown', 'systemctl', 'service', 'wget', 'tar',
                'unzip', 'ollama', 'nvcc', 'jetson', 'jtop', 'nvidia',
                'make', 'cmake', './', 'bash', 'sh ', 'tee', 'touch',
                'find', 'grep', 'sed', 'awk', 'xargs', 'head', 'tail',
            ))
        )

        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8)
        if is_comment:
            r.font.color.rgb = hex_rgb(CODE_COMMENT)
            r.italic = True  # obs#7: italic for all comments
        elif is_output:
            r.font.color.rgb = hex_rgb(CODE_OUTPUT)
            r.italic = True
        else:
            r.font.color.rgb = hex_rgb(CODE_FG)


def add_code_block(doc, code, lang='bash'):
    """obs#7: splits at separator lines, renders each section as a separate table."""
    sections = split_code_at_separators(code)
    for idx, section in enumerate(sections):
        _render_code_table(doc, section, lang)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(2 if idx < len(sections) - 1 else 6)


def add_callout(doc, text, kind='NOTA'):
    kind_upper = kind.upper().strip().replace('ATENCION', 'ATENCIÓN')
    labels = {
        'IMPORTANTE': ('IMPORTANTE', CALLOUT_IMP),
        'NOTA':       ('NOTA',       CALLOUT_NOTE),
        'CONSEJO':    ('CONSEJO',    CALLOUT_TIP),
        'ADVERTENCIA':('ADVERTENCIA',CALLOUT_WARN),
        'ATENCIÓN':   ('ATENCIÓN',   CALLOUT_WARN),
    }
    label, border_color = labels.get(kind_upper, ('NOTA', CALLOUT_NOTE))

    tbl = doc.add_table(rows=1, cols=1)
    table_full_width(tbl)
    cell = tbl.cell(0, 0)
    cell_bg(cell, 'F8F8F8')
    set_full_border(cell, border_color, sz='4', left_sz='16')
    set_cell_padding(cell, top=100, left=200, bottom=100, right=160)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    r_label = p.add_run(f'{label}: ')
    r_label.bold = True
    r_label.font.name = 'Amazon Endure' if _has_endure_font() else 'Arial'
    r_label.font.size = Pt(9)
    r_label.font.color.rgb = hex_rgb(border_color)

    # Texto del callout con inline formatting
    _add_inline_formatted_runs(p, text.strip())

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(6)


def _has_endure_font():
    return True  # La plantilla ya embebe la fuente


def add_table(doc, headers, rows, description=None):
    """obs#6: white borders 1.5pt, EAEAEA uniform bg, repeat header, center/justify by length."""
    if description:
        p_desc = doc.add_paragraph()
        r_desc = p_desc.add_run(description)
        r_desc.font.size = Pt(8)   # obs#11
        r_desc.italic = True
        p_desc.paragraph_format.space_before = Pt(5)
        p_desc.paragraph_format.space_after = Pt(2)

    n_cols = max(len(headers), 1)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    try:
        tbl.style = 'Table Grid'
    except KeyError:
        pass
    table_full_width(tbl)

    # ── Fila de encabezado ────────────────────────────────────────────────────
    hdr_row = tbl.rows[0]
    set_row_height(hdr_row, 0.8)     # obs#6: 0.8cm mínimo
    set_repeat_header(hdr_row)       # obs#6: repetir en cada página
    for i, h in enumerate(headers[:n_cols]):
        cell = hdr_row.cells[i]
        cell_bg(cell, NVIDIA_TABLE)
        set_all_cell_borders(cell, TABLE_BORDER, TABLE_BORDER_SZ)
        set_cell_valign(cell)
        set_cell_padding(cell, top=60, left=120, bottom=60, right=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        h_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', str(h))
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                       if len(h_clean.replace(' ', '')) < 30
                       else WD_ALIGN_PARAGRAPH.JUSTIFY)
        r = p.add_run(h_clean)
        r.bold = True
        r.font.size = Pt(8)     # obs#5
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Filas de datos ─────────────────────────────────────────────────────────
    for ri, row in enumerate(rows):
        tbl_row = tbl.rows[ri + 1]
        set_row_height(tbl_row, 0.5)  # obs#6: 0.5cm mínimo
        for ci, val in enumerate(row[:n_cols]):
            cell = tbl_row.cells[ci]
            cell_bg(cell, TABLE_BG)   # obs#6: uniforme EAEAEA, sin alternancia
            set_all_cell_borders(cell, TABLE_BORDER, TABLE_BORDER_SZ)
            set_cell_valign(cell)
            set_cell_padding(cell, top=40, left=120, bottom=40, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            val_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', str(val))
            val_clean = re.sub(r'\*(.+?)\*', r'\1', val_clean)
            val_clean = re.sub(r'`(.+?)`', r'\1', val_clean)
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                           if len(val_clean.replace(' ', '')) < 30
                           else WD_ALIGN_PARAGRAPH.JUSTIFY)
            r = p.add_run(val_clean)
            r.font.size = Pt(8)  # obs#5

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(6)


def add_bullets(doc, items, numbered=False, is_front_matter=False):
    style = get_style(doc, STYLE_BODY, 'Normal')
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        prefix = f'{i}. ' if numbered else '• '
        r_prefix = p.add_run(prefix)
        r_prefix.font.size = Pt(10)  # obs#5
        _add_inline_formatted_runs(p, item.strip(), font_size=Pt(10))  # obs#5
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)


def add_image_caption(doc, caption_text):
    p = doc.add_paragraph(style=get_style(doc, STYLE_COPYRIGHT, 'Normal'))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption_text)
    r.italic = True
    r.font.size = Pt(8)  # obs#5+#11


# ═══════════════════════════════════════════════════════════════════════════════
# GENERACIÓN DE INFOGRAFÍAS
# ═══════════════════════════════════════════════════════════════════════════════

C_DARK   = '#0F3D3D'
C_ACCENT = '#1D9CB8'
C_MID    = '#1A5A6A'
C_LIGHT  = '#E8F4F8'
C_WHITE  = '#FFFFFF'
C_GRAY   = '#CCCCCC'
C_TEXT   = '#F0F0F0'
C_DARK_TEXT = '#1A1A1A'


def fig_setup(w=7.0, h=4.5):   # obs#3: increased default height from 3.8 to 4.5
    fig, ax = plt.subplots(figsize=(w, h))
    fig.patch.set_facecolor(C_DARK)
    ax.set_facecolor(C_DARK)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis('off')
    return fig, ax


def draw_box(ax, x, y, w, h, text, sub='', color=C_MID, text_color=C_WHITE, fontsize=9, subsize=8.0):
    rect = mpatches.FancyBboxPatch((x, y), w, h,
                                    boxstyle='round,pad=0.1',
                                    facecolor=color, edgecolor=C_ACCENT, linewidth=1.2)
    ax.add_patch(rect)
    cy = y + h / 2
    if sub:
        ax.text(x + w/2, cy + 0.22, text, ha='center', va='center',   # obs#3: was 0.18
                color=text_color, fontsize=fontsize, fontweight='bold')
        ax.text(x + w/2, cy - 0.38, sub, ha='center', va='center',    # obs#3: was -0.28
                color=C_GRAY, fontsize=subsize)
    else:
        ax.text(x + w/2, cy, text, ha='center', va='center',
                color=text_color, fontsize=fontsize, fontweight='bold',
                wrap=True)


def draw_arrow(ax, x1, y, x2, color=C_ACCENT):
    if isinstance(color, (int, float)):
        color = C_ACCENT  # called with y passed twice by mistake — ignore
    ax.annotate('', xy=(x2, y), xytext=(x1, y),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.5))


def draw_arrow_v(ax, x, y1, y2, color=C_ACCENT):
    ax.annotate('', xy=(x, y2), xytext=(x, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=1.5))


def fig_title(ax, title, y=9.3):
    ax.text(5, y, title, ha='center', va='center',
            color=C_ACCENT, fontsize=10, fontweight='bold')


def save_fig(fig, path):
    INFOGRAFIAS_DIR.mkdir(exist_ok=True)
    fig.savefig(str(path), dpi=150, bbox_inches='tight',
                facecolor=C_DARK, edgecolor='none')
    plt.close(fig)


# ── Infografía 1 — Cap 1: Arquitectura del sistema al final del libro ─────────
def infografia_cap01_arquitectura_sistema():
    fig, ax = fig_setup(7.0, 4.8)
    fig_title(ax, 'Arquitectura del Sistema al Final del Libro')
    layers = [
        (C_MID,    'INTERNET / CLIENTES',   ':80/:443'),
        ('#0A7A7A', 'Cloudflare Tunnel + Nginx', 'Reverse proxy + SSL'),
        (C_MID,    'Stack IA Agéntica',     'OpenClaw · NemoClaw · Open WebUI'),
        ('#0A5050', 'Motores de Inferencia', 'Ollama · llama.cpp · vLLM'),
        ('#083C3C', 'Hardware Jetson',       'GPU Ampere sm_87 · 64 GB RAM Unificada'),
    ]
    total = len(layers)
    h = 1.4   # obs#3: was 1.3
    gap = 0.25  # obs#3: was 0.18
    start_y = 8.8 - (total * (h + gap))
    for i, (color, label, sub) in enumerate(layers):
        y = start_y + i * (h + gap)
        draw_box(ax, 0.3, y, 9.4, h, label, sub, color=color, fontsize=9)
        if i < total - 1:
            draw_arrow_v(ax, 5.0, y - gap, y, color=C_ACCENT)
    path = INFOGRAFIAS_DIR / 'fig-cap01-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 2 — Cap 2: Force Recovery ─────────────────────────────────────
def infografia_cap02_force_recovery():
    fig, ax = fig_setup(7.0, 3.6)
    fig_title(ax, 'Pasos: Modo Force Recovery en Jetson AGX Orin')
    steps = [
        ('1', 'Apagar\nel Jetson'),
        ('2', 'Conectar\nUSB-C a PC'),
        ('3', 'Mantener\nREC'),
        ('4', 'Pulsar\nRESET'),
        ('5', 'Soltar\nREC'),
        ('6', 'Verificar\nlsusb'),
    ]
    n = len(steps)
    bw = 1.3
    gap = (9.4 - n * bw) / (n - 1)
    y_box = 3.5
    for i, (num, label) in enumerate(steps):
        x = 0.3 + i * (bw + gap)
        draw_box(ax, x, y_box, bw, 3.0, label, sub=num,
                 color=C_MID if i % 2 == 0 else '#0A7A7A', fontsize=8)
        if i < n - 1:
            draw_arrow(ax, x + bw, y_box + 1.5, x + bw + gap, y_box + 1.5)
    path = INFOGRAFIAS_DIR / 'fig-cap02-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 3 — Cap 4: Memoria Unificada ──────────────────────────────────
def infografia_cap04_memoria_unificada():
    fig, ax = fig_setup(7.0, 4.0)
    fig_title(ax, 'Arquitectura de Memoria Unificada del Jetson AGX Orin 64GB')
    # Bloque grande de RAM
    rect_main = mpatches.FancyBboxPatch((0.3, 1.5), 9.4, 6.0,
                                         boxstyle='round,pad=0.15',
                                         facecolor='#0A3030', edgecolor=C_ACCENT, linewidth=2)
    ax.add_patch(rect_main)
    ax.text(5, 8.2, '64 GB RAM Unificada LPDDR5X', ha='center', color=C_ACCENT,
            fontsize=10, fontweight='bold')
    # Subdivisiones
    slots = [
        (0.5, 5.5, 3.0, 1.6, 'CPU (12 núcleos\nArm Cortex-A78AE)', '~12 GB OS'),
        (3.8, 5.5, 5.7, 1.6, 'GPU (2048 CUDA cores\nAmpere sm_87)', 'Acceso directo sin copia'),
        (0.5, 3.5, 3.0, 1.6, 'Ollama 7B Model', '~18 GB'),
        (3.8, 3.5, 2.6, 1.6, 'vLLM 13B', '~28 GB'),
        (6.7, 3.5, 2.8, 1.6, 'SD WebUI SDXL', '~12 GB'),
        (0.5, 1.7, 9.0, 1.5, 'Sistema Operativo Ubuntu 24.04 (reposo ~12 GB)', ''),
    ]
    for (x, y, w, h, label, sub) in slots:
        color = '#0A5A5A' if 'GPU' in label else '#0A4A4A' if 'CPU' in label else C_MID
        draw_box(ax, x, y, w, h, label, sub, color=color, fontsize=7.5, subsize=7)
    path = INFOGRAFIAS_DIR / 'fig-cap04-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 4 — Cap 5: nvpmodel vs jetson_clocks ─────────────────────────
def infografia_cap05_rendimiento():
    fig, ax = fig_setup(7.0, 3.8)
    fig_title(ax, 'Dos Mecanismos de Control de Rendimiento')
    # nvpmodel
    draw_box(ax, 0.3, 2.0, 4.3, 5.5, 'nvpmodel', sub='Perfiles de potencia predefinidos\n(MAXN, 15W, 30W, 50W)', color=C_MID, fontsize=9)
    items_l = ['Controla CPU, GPU y DLA', 'Define límites de TDP', 'Persistente entre reinicios', 'sudo nvpmodel -m 0 (MAXN)']
    for i, t in enumerate(items_l):
        ax.text(0.6, 6.8 - i * 0.9, f'• {t}', color=C_TEXT, fontsize=7.5)
    # jetson_clocks
    draw_box(ax, 5.4, 2.0, 4.3, 5.5, 'jetson_clocks', sub='Fija frecuencias al máximo\nabsoluto inmediatamente', color='#0A7A7A', fontsize=9)
    items_r = ['Override manual del gobernador', 'Máx. latencia de inferencia', 'No persiste sin servicio', 'sudo jetson_clocks --show']
    for i, t in enumerate(items_r):
        ax.text(5.7, 6.8 - i * 0.9, f'• {t}', color=C_TEXT, fontsize=7.5)
    ax.text(4.85, 4.5, '⬡', ha='center', color=C_ACCENT, fontsize=14)
    ax.text(4.85, 3.5, 'usar\njuntos', ha='center', color=C_GRAY, fontsize=7)
    path = INFOGRAFIAS_DIR / 'fig-cap05-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 5 — Cap 9: GPU en Contenedores ────────────────────────────────
def infografia_cap09_gpu_containers():
    fig, ax = fig_setup(7.0, 4.8)
    fig_title(ax, 'Arquitectura de GPU en Contenedores Jetson')
    layers = [
        (C_MID,    'Aplicación / Modelo IA', '(Python, PyTorch, TensorRT)'),
        ('#0A6A6A', 'CUDA Runtime 13.2.1', '(dentro del contenedor)'),
        ('#0A5050', 'NVIDIA Container Toolkit', '(nvidia-ctk, libnvidia-container)'),
        ('#0A3A3A', 'Driver NVIDIA L4T', '(Host: /dev/nvhost-*, /dev/tegra-*)'),
        ('#083030', 'Hardware GPU Ampere sm_87', '(2048 CUDA cores · 64 GB RAM Unificada)'),
    ]
    h, gap = 1.3, 0.25  # obs#3: was 1.1, 0.2
    start_y = 1.3
    for i, (color, label, sub) in enumerate(layers):
        y = start_y + i * (h + gap)
        draw_box(ax, 0.4, y, 9.2, h, label, sub, color=color, fontsize=8.5)
        if i < len(layers) - 1:
            draw_arrow_v(ax, 5.0, y - gap + 0.05, y - 0.05)
    ax.text(9.8, 4.0, 'Docker\nContainer', ha='center', color=C_ACCENT, fontsize=7.5, style='italic')
    ax.annotate('', xy=(9.2, 5.8), xytext=(9.8, 5.0),
                arrowprops=dict(arrowstyle='->', color=C_ACCENT, lw=1))
    path = INFOGRAFIAS_DIR / 'fig-cap09-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 6 — Cap 11: Arquitectura OpenClaw ─────────────────────────────
def infografia_cap11_openclaw():
    fig, ax = fig_setup(7.0, 4.0)
    fig_title(ax, 'Arquitectura del Stack Agéntico en Jetson AGX Orin')
    # Entradas
    entradas = ['Telegram\nBot', 'Web UI\n:3000', 'API REST\n:18789']
    for i, e in enumerate(entradas):
        draw_box(ax, 0.3 + i * 2.0, 7.2, 1.7, 1.3, e, color='#0A5050', fontsize=8)
        draw_arrow(ax, 0.3 + i * 2.0 + 1.7, 7.85, 6.5, 7.85)
    # OpenClaw
    draw_box(ax, 6.5, 6.5, 3.2, 2.0, 'OpenClaw\nGateway', sub=':18789', color=C_MID, fontsize=9)
    draw_arrow_v(ax, 8.1, 6.5, 5.0)
    # NemoClaw
    draw_box(ax, 6.5, 4.0, 3.2, 2.0, 'NemoClaw\nSecurity L7', sub='JWT + Rate Limit', color='#0A7A7A', fontsize=9)
    draw_arrow_v(ax, 8.1, 4.0, 2.5)
    # Backends
    backends = [('vLLM\n:8000', C_MID), ('llama.cpp\n:8080', '#0A5050'), ('Ollama\n:11434', '#0A6060')]
    for i, (label, color) in enumerate(backends):
        draw_box(ax, 1.0 + i * 3.0, 1.0, 2.5, 1.6, label, color=color, fontsize=8)
        draw_arrow_v(ax, 8.1, 2.5, 2.0 - 0.2)
    path = INFOGRAFIAS_DIR / 'fig-cap11-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 7 — Cap 11A: OpenClaw JetPack 7.2 ─────────────────────────────
def infografia_cap11a_openclaw():
    fig, ax = fig_setup(7.0, 3.8)
    fig_title(ax, 'OpenClaw — Bot de Telegram con IA Local (JetPack 7.2)')
    # Flujo
    nodes = [
        ('Usuario\nTelegram', C_MID),
        ('Telegram\nAPI Cloud', '#0A5050'),
        ('OpenClaw\n:18789', C_MID),
        ('Model\nBackend', '#0A7A7A'),
        ('Respuesta\nal usuario', C_MID),
    ]
    bw, gap = 1.6, 0.3
    total_w = len(nodes) * bw + (len(nodes) - 1) * gap
    start_x = (10 - total_w) / 2
    y_box = 3.8
    for i, (label, color) in enumerate(nodes):
        x = start_x + i * (bw + gap)
        draw_box(ax, x, y_box, bw, 2.5, label, color=color, fontsize=8)
        if i < len(nodes) - 1:
            draw_arrow(ax, x + bw, y_box + 1.25, x + bw + gap, y_box + 1.25)
    # Config files
    draw_box(ax, 0.3, 1.0, 4.0, 1.8, '.env / config.yaml', sub='API_KEY · MODEL · BACKEND_URL', color='#083030', fontsize=8)
    draw_box(ax, 5.7, 1.0, 4.0, 1.8, 'Servicios del Jetson', sub='vLLM · Ollama · llama.cpp', color='#083030', fontsize=8)
    path = INFOGRAFIAS_DIR / 'fig-cap11a-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 8 — Cap 11B: NemoClaw ─────────────────────────────────────────
def infografia_cap11b_nemoclaw():
    fig, ax = fig_setup(7.0, 4.8)
    fig_title(ax, 'NemoClaw — Capa de Seguridad L7 sobre OpenClaw')
    layers = [
        ('Cliente Externo',              '#083030', '(Telegram / Web / API)'),
        ('Cloudflare Tunnel + SSL',      '#0A3A3A', '(TLS 1.3 · DDoS mitigation)'),
        ('NemoClaw Security Middleware', C_MID,     '(JWT · Rate Limiting · Input Sanitization)'),
        ('OpenClaw Gateway :18789',      '#0A5A5A', '(Routing · Context management)'),
        ('LLM Backend (vLLM / Ollama)',  '#0A7070', '(Inferencia local · 64 GB RAM)'),
    ]
    h, gap = 1.2, 0.28  # obs#3: was 1.0, 0.25
    start_y = 1.5  # obs#3: recalculated
    for i, (label, color, sub) in enumerate(layers):
        y = start_y + i * (h + gap)
        draw_box(ax, 0.4, y, 9.2, h, label, sub, color=color, fontsize=8)
        if i < len(layers) - 1:
            lcolor = C_ACCENT if i == 2 else '#555555'
            draw_arrow_v(ax, 5.0, y - gap + 0.05, y - 0.05, color=lcolor)
    ax.text(9.9, 4.8, 'BLOQUEADO\nsin JWT', ha='center', color=CALLOUT_WARN, fontsize=7, fontweight='bold')
    path = INFOGRAFIAS_DIR / 'fig-cap11b-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 9 — Cap 13: Pipeline STT→LLM→TTS ─────────────────────────────
def infografia_cap13_stt_llm_tts():
    fig, ax = fig_setup(7.0, 4.0)
    fig_title(ax, 'Pipeline STT → LLM → TTS en Jetson AGX Orin')
    nodes = [
        ('Micrófono\nUSB', 'pyaudio\n(VAD)', '~1s'),
        ('faster-\nwhisper', 'STT → texto', '~0.5s'),
        ('vLLM\nQwen3.5-4B', 'LLM → respuesta', '~1–2s'),
        ('piper-\ntts', 'TTS → WAV', '<200ms'),
        ('Altavoz\nUSB', 'Audio output', '—'),
    ]
    bw, gap = 1.5, 0.3
    total_w = len(nodes) * bw + (len(nodes) - 1) * gap
    start_x = (10 - total_w) / 2
    y_main = 5.5
    for i, (label, sub, latency) in enumerate(nodes):
        x = start_x + i * (bw + gap)
        color = C_MID if i % 2 == 0 else '#0A7070'
        draw_box(ax, x, y_main, bw, 2.8, label, sub=sub, color=color, fontsize=8)
        ax.text(x + bw/2, y_main - 0.7, latency, ha='center', color=C_ACCENT, fontsize=7.5)
        if i < len(nodes) - 1:
            draw_arrow(ax, x + bw, y_main + 1.4, x + bw + gap, y_main + 1.4)
    # Latencia total
    draw_box(ax, 1.0, 1.0, 8.0, 1.8, 'Latencia Total: <3 segundos', sub='STT 0.5s + LLM 1–2s + TTS <0.2s', color='#083030', fontsize=9)
    path = INFOGRAFIAS_DIR / 'fig-cap13-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 10 — Cap 14: Arquitectura Imagen/Video ────────────────────────
def infografia_cap14_imagen_video():
    fig, ax = fig_setup(7.0, 4.0)
    fig_title(ax, 'Arquitectura de Generación de Imágenes — Modelos Compartidos')
    # Modelos compartidos
    draw_box(ax, 2.5, 7.2, 5.0, 1.6, 'Modelos Compartidos', sub='/data/models/stable-diffusion/', color='#083030', fontsize=9)
    # ComfyUI y SD WebUI
    draw_box(ax, 0.4, 4.5, 4.2, 2.3, 'ComfyUI\n:7860', sub='Flujos de trabajo avanzados\nAnimateDiff · ControlNet', color=C_MID, fontsize=8.5)
    draw_box(ax, 5.4, 4.5, 4.2, 2.3, 'SD WebUI\n:7861', sub='Interfaz clásica\nExtensiones · Scripts', color='#0A7070', fontsize=8.5)
    draw_arrow_v(ax, 2.5, 7.2, 6.8)
    draw_arrow_v(ax, 7.5, 7.2, 6.8)
    # GPU
    draw_box(ax, 0.4, 1.5, 9.2, 2.5, 'GPU Ampere sm_87 · CUDA 13.2.1 · 64 GB RAM Unificada',
             sub='TensorRT · CuDNN · Aceleración nativa para SDXL y AnimateDiff', color='#0A3030', fontsize=8)
    draw_arrow_v(ax, 2.5, 4.5, 4.0)
    draw_arrow_v(ax, 7.5, 4.5, 4.0)
    path = INFOGRAFIAS_DIR / 'fig-cap14-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 11 — Cap 15: N8N Webhook→LLM→Email ───────────────────────────
def infografia_cap15_pipeline_webhook():
    fig, ax = fig_setup(7.0, 3.6)
    fig_title(ax, 'N8N Pipeline: Webhook → LLM → Email')
    nodes = [
        ('Webhook\nTrigger', ':5678/webhook'),
        ('HTTP Request\n→ OpenClaw', 'POST :18789'),
        ('LLM\nQwen3.5', 'Genera respuesta'),
        ('Gmail /\nSMTP Node', 'Envía email'),
        ('Destinatario\nFinal', 'Respuesta IA'),
    ]
    bw, gap = 1.5, 0.28
    start_x = 0.3
    y_box = 3.5
    for i, (label, sub) in enumerate(nodes):
        color = C_MID if i % 2 == 0 else '#0A6060'
        draw_box(ax, start_x + i * (bw + gap), y_box, bw, 3.2, label, sub=sub, color=color, fontsize=8)
        if i < len(nodes) - 1:
            draw_arrow(ax, start_x + i * (bw + gap) + bw, y_box + 1.6,
                       start_x + (i + 1) * (bw + gap), y_box + 1.6)
    path = INFOGRAFIAS_DIR / 'fig-cap15-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 12 — Cap 15: Email→LLM→Respuesta ─────────────────────────────
def infografia_cap15_pipeline_email():
    fig, ax = fig_setup(7.0, 3.6)
    fig_title(ax, 'N8N Pipeline: Email Entrante → LLM → Respuesta Automática')
    nodes = [
        ('IMAP\nTrigger', 'Nuevo email\ndetectado'),
        ('Extraer\nContenido', 'Subject + Body'),
        ('OpenClaw\nLLM', 'Genera borrador'),
        ('Revisión\nCondicional', 'Score > 0.8?'),
        ('Enviar\nRespuesta', 'Gmail SMTP'),
    ]
    bw, gap = 1.5, 0.28
    start_x = 0.3
    y_box = 3.5
    for i, (label, sub) in enumerate(nodes):
        color = '#0A7070' if i == 3 else C_MID
        draw_box(ax, start_x + i * (bw + gap), y_box, bw, 3.2, label, sub=sub, color=color, fontsize=7.5)
        if i < len(nodes) - 1:
            draw_arrow(ax, start_x + i * (bw + gap) + bw, y_box + 1.6,
                       start_x + (i + 1) * (bw + gap), y_box + 1.6)
    # Loop de revisión
    ax.annotate('', xy=(start_x + 2 * (bw + gap) + bw / 2, y_box + 3.2 + 0.3),
                xytext=(start_x + 3 * (bw + gap) + bw / 2, y_box + 3.2 + 0.3),
                arrowprops=dict(arrowstyle='->', color=CALLOUT_WARN, lw=1.2,
                                connectionstyle='arc3,rad=-0.3'))
    ax.text(4.5, y_box + 4.0, 'Score < 0.8\n→ regenerar', ha='center', color=CALLOUT_WARN, fontsize=7)
    path = INFOGRAFIAS_DIR / 'fig-cap15-02.png'
    save_fig(fig, path)
    return path


# ── Infografía 13 — Cap 15: N8N+OpenClaw+vLLM ───────────────────────────────
def infografia_cap15_integracion():
    fig, ax = fig_setup(7.0, 4.2)
    fig_title(ax, 'Integración: N8N + OpenClaw + vLLM')
    draw_box(ax, 0.3, 6.5, 2.8, 2.5, 'N8N\n:5678', sub='Orquestador\nde flujos', color=C_MID, fontsize=9)
    draw_box(ax, 3.7, 6.5, 2.8, 2.5, 'OpenClaw\n:18789', sub='Gateway IA\n+ NemoClaw', color='#0A7070', fontsize=9)
    draw_box(ax, 7.1, 6.5, 2.6, 2.5, 'vLLM\n:8000', sub='Qwen3.5-35B\n35B params', color='#0A5050', fontsize=9)
    draw_arrow(ax, 3.1, 7.75, 3.7, 7.75)
    draw_arrow(ax, 6.5, 7.75, 7.1, 7.75)
    # Fuentes de trigger
    triggers = ['Webhook', 'Cron Job', 'Gmail', 'Telegram']
    for i, t in enumerate(triggers):
        x = 0.5 + i * 2.2
        draw_box(ax, x, 3.8, 1.9, 1.5, t, color='#083030', fontsize=7.5)
        draw_arrow_v(ax, x + 0.95, 5.3, 6.5)
    # Destinos
    draw_box(ax, 0.3, 1.0, 9.4, 1.8, 'Resultados: Email · Slack · DB · Telegram · Webhook de respuesta',
             color='#083030', fontsize=8)
    draw_arrow_v(ax, 5.0, 6.5, 2.8)
    path = INFOGRAFIAS_DIR / 'fig-cap15-03.png'
    save_fig(fig, path)
    return path


# ── Infografía 14 — Cap 17: Estado objetivo boot limpio ──────────────────────
def infografia_cap17_boot_state():
    fig, ax = fig_setup(7.0, 4.0)
    fig_title(ax, 'Estado Objetivo del Sistema tras Boot Limpio')
    services = [
        ('systemd-resolved', 'DNS-over-TLS activo', '✓'),
        ('ufw (firewall)', 'Reglas cargadas', '✓'),
        ('docker.service', 'NVIDIA runtime listo', '✓'),
        ('ollama.service', 'API :11434 respondiendo', '✓'),
        ('watchdog.service', 'Monitor de OOM activo', '✓'),
        ('swap / ZRAM', '16 GB ZRAM + 32 GB NVMe swap', '✓'),
    ]
    h = 0.98
    start_y = 1.8
    for i, (svc, desc, status) in enumerate(services):
        y = start_y + i * h
        color = C_MID if i % 2 == 0 else '#0A5050'
        draw_box(ax, 0.4, y, 6.8, h - 0.05, svc, sub=desc, color=color, fontsize=8)
        ax.text(8.0, y + h / 2, status, ha='center', va='center',
                color='#00CC77', fontsize=12, fontweight='bold')
    ax.text(8.0, 1.3, 'Estado', ha='center', color=C_ACCENT, fontsize=7.5)
    path = INFOGRAFIAS_DIR / 'fig-cap17-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 15 — Cap 19: Ciclo de desarrollo en Jetson ────────────────────
def infografia_cap19_ciclo_desarrollo():
    fig, ax = fig_setup(7.0, 4.5)
    fig_title(ax, 'Patrón de Desarrollo en Jetson — Ciclo de 5 Pasos')
    import numpy as np
    steps = [
        ('1\nVerificar\nmemoria', 30),
        ('2\nLanzar\nmotor IA', 102),
        ('3\nDesarrollar\ny testear', 174),
        ('4\nTerminar\ny limpiar', 246),
        ('5\nVerificar\nmemoria libre', 318),
    ]
    cx, cy, r = 5.0, 4.5, 2.8
    colors = [C_MID, '#0A7070', '#0A5050', '#0A6060', '#0A7A7A']
    for i, (label, angle_deg) in enumerate(steps):
        angle = np.radians(angle_deg)
        x = cx + r * np.cos(angle)
        y = cy + r * np.sin(angle)
        draw_box(ax, x - 0.9, y - 0.65, 1.8, 1.3, label, color=colors[i], fontsize=7.5)
        # Flecha al siguiente
        next_angle = np.radians(steps[(i + 1) % len(steps)][1])
        nx = cx + r * np.cos(next_angle)
        ny = cy + r * np.sin(next_angle)
        ax.annotate('', xy=(nx, ny), xytext=(x, y),
                    arrowprops=dict(arrowstyle='->', color=C_ACCENT, lw=1.2,
                                    connectionstyle='arc3,rad=0.3'))
    ax.text(cx, cy, 'Jetson\nAGX Orin', ha='center', va='center',
            color=C_ACCENT, fontsize=9, fontweight='bold')
    path = INFOGRAFIAS_DIR / 'fig-cap19-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 16 — Cap 20: jetson-containers vs Docker Hub ──────────────────
def infografia_cap20_jetson_containers():
    fig, ax = fig_setup(7.0, 3.8)
    fig_title(ax, 'jetson-containers vs Docker Hub Genérico')
    col_titles = ['Docker Hub genérico', 'jetson-containers (dustynv)']
    col_colors = ['#4A1A1A', C_MID]
    pros_l = ['Imágenes x86_64\n(incompatibles con ARM)', 'Sin soporte CUDA Jetson', 'Sin optimizaciones L4T', 'Requiere compilar todo\ndesde cero']
    pros_r = ['Imágenes nativas ARM64', 'CUDA 13.2.1 + TensorRT', 'Optimizadas para L4T r39.2', 'Listo para usar en minutos']
    for col, (title, color, items) in enumerate(zip(col_titles, col_colors, [pros_l, pros_r])):
        x_base = 0.4 + col * 5.0
        draw_box(ax, x_base, 6.8, 4.5, 1.4, title, color=color, fontsize=9)
        for i, item in enumerate(items):
            icon = '✗' if col == 0 else '✓'
            icon_color = CALLOUT_WARN if col == 0 else '#00CC77'
            ax.text(x_base + 0.3, 5.8 - i * 1.2, icon, color=icon_color, fontsize=10, va='top')
            ax.text(x_base + 0.7, 5.8 - i * 1.2, item, color=C_TEXT, fontsize=7.5, va='top')
    path = INFOGRAFIAS_DIR / 'fig-cap20-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 17 — Cap 20: Mapa de puertos ──────────────────────────────────
def infografia_cap20_puertos():
    fig, ax = fig_setup(7.0, 4.5)
    fig_title(ax, 'Mapa de Puertos — Todos los Servicios del Jetson')
    services = [
        ('Ollama',      ':11434', 'Inferencia LLM'),
        ('vLLM',        ':8000',  'Inferencia LLM (OpenAI compat.)'),
        ('llama.cpp',   ':8080',  'Inferencia LLM'),
        ('OpenClaw',    ':18789', 'Gateway agéntico'),
        ('Open WebUI',  ':3000',  'Interfaz gráfica web'),
        ('ComfyUI',     ':7860',  'Generación de imágenes'),
        ('SD WebUI',    ':7861',  'Generación de imágenes'),
        ('N8N',         ':5678',  'Automatización de flujos'),
        ('Nginx/SAAS',  ':80/:443','Exposición a Internet'),
    ]
    h = 0.78
    start_y = 1.5
    for i, (svc, port, desc) in enumerate(services):
        y = start_y + i * h
        color = C_MID if i % 2 == 0 else '#0A5050'
        draw_box(ax, 0.3, y, 3.2, h - 0.04, svc, color=color, fontsize=8)
        draw_box(ax, 3.7, y, 1.8, h - 0.04, port, color='#0A3A3A', fontsize=8)
        ax.text(5.8, y + h / 2, desc, va='center', color=C_TEXT, fontsize=7.5)
    ax.text(1.9, 1.2, 'Servicio', ha='center', color=C_ACCENT, fontsize=7.5, fontweight='bold')
    ax.text(4.6, 1.2, 'Puerto', ha='center', color=C_ACCENT, fontsize=7.5, fontweight='bold')
    path = INFOGRAFIAS_DIR / 'fig-cap20-02.png'
    save_fig(fig, path)
    return path


# ── Infografía 18 — Cap 23: Agencia de Turismo Virtual ───────────────────────
def infografia_cap23_turismo():
    fig, ax = fig_setup(7.0, 4.2)
    fig_title(ax, 'Arquitectura — Agencia de Turismo Virtual con OpenClaw/NemoClaw')
    draw_box(ax, 3.5, 7.8, 3.0, 1.4, 'Telegram Bot', sub='Usuario hace consulta', color='#083030', fontsize=8)
    draw_arrow_v(ax, 5.0, 7.8, 7.2)
    draw_box(ax, 2.5, 5.5, 5.0, 1.7, 'OpenClaw + NemoClaw', sub='Routing · JWT · Rate Limiting', color=C_MID, fontsize=9)
    draw_arrow_v(ax, 5.0, 5.5, 4.8)
    # Agentes
    agents = [('Agente\nHoteles', 0.4), ('Agente\nVuelos', 3.0), ('Agente\nActividades', 5.6), ('Agente\nPresupuesto', 8.2)]
    for label, x in agents:
        draw_box(ax, x, 2.5, 1.5, 2.3, label, color='#0A6060', fontsize=7.5)
        draw_arrow_v(ax, x + 0.75, 4.8, 4.8)
    draw_box(ax, 0.3, 0.8, 9.4, 1.5, 'vLLM Qwen3.5-35B · ChromaDB RAG · Datos de Turismo',
             color='#0A3030', fontsize=8)
    for label, x in agents:
        draw_arrow_v(ax, x + 0.75, 2.5, 2.3)
    path = INFOGRAFIAS_DIR / 'fig-cap23-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 19 — Cap 28: Arquitectura AIaaS ───────────────────────────────
def infografia_cap28_aiaas():
    fig, ax = fig_setup(7.0, 5.0)
    fig_title(ax, 'Arquitectura AIaaS — Jetson Expuesto a Internet')
    layers = [
        ('Internet / Clientes',            '#083030', ':443 HTTPS'),
        ('Cloudflare Tunnel',              '#0A3A3A', 'Sin IP pública · DDoS · TLS automático'),
        ('Nginx Reverse Proxy',            C_MID,     ':80 → servicios internos'),
        ('JWT Middleware (NemoClaw)',       '#0A6060', 'Autenticación · Rate Limiting'),
        ('Servicios IA Internos',          '#0A5050', 'vLLM :8000 · STT · TTS · OpenClaw :18789'),
        ('Hardware Jetson AGX Orin 64GB',  '#0A3030', 'GPU Ampere · Inferencia local'),
    ]
    h, gap = 1.2, 0.25  # obs#3: was 1.05, 0.18
    start_y = 1.2  # obs#3: recalculated
    for i, (label, color, sub) in enumerate(layers):
        y = start_y + i * (h + gap)
        draw_box(ax, 0.4, y, 9.2, h, label, sub, color=color, fontsize=8.5)
        if i < len(layers) - 1:
            lw = 2.0 if i == 2 else 1.2
            draw_arrow_v(ax, 5.0, y - gap + 0.05, y - 0.05,
                         color=C_ACCENT if i == 2 else '#555555')
    path = INFOGRAFIAS_DIR / 'fig-cap28-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 20 — Capstone 01: Agencia de IA con presencia web ──────────────
def infografia_capstone01():
    fig, ax = fig_setup(7.0, 4.5)
    fig_title(ax, 'Capstone 01 — Agencia de IA con Presencia Web (RAM por componente)')
    components = [
        ('Ubuntu 24.04 (OS)',        12, '#0A3030'),
        ('Open WebUI :3000',          3, '#0A4A4A'),
        ('N8N :5678',                 2, '#0A4A4A'),
        ('NemoClaw + Nginx',          1, '#0A5050'),
        ('OpenClaw :18789',           2, C_MID),
        ('vLLM Qwen3.5-35B',         30, '#0A7070'),
        ('Buffer / SO overhead',       6, '#083030'),
    ]
    total = 64
    y_base = 1.5
    bar_h = 0.75
    x_scale = 8.5 / total
    y_cur = y_base
    for label, gb, color in components:
        w = gb * x_scale
        rect = mpatches.FancyBboxPatch((0.5, y_cur), w, bar_h - 0.05,
                                        boxstyle='round,pad=0.05',
                                        facecolor=color, edgecolor=C_ACCENT, linewidth=0.8)
        ax.add_patch(rect)
        ax.text(0.5 + w + 0.1, y_cur + bar_h / 2, f'{label} ({gb}GB)',
                va='center', color=C_TEXT, fontsize=7.5)
        y_cur += bar_h
    # Línea total
    ax.axvline(x=0.5 + total * x_scale, color=CALLOUT_WARN, lw=1.5, linestyle='--')
    ax.text(0.5 + total * x_scale, y_cur + 0.3, '64 GB\ntotal', ha='center',
            color=CALLOUT_WARN, fontsize=7.5)
    ax.text(5.0, 1.0, 'Pico de uso: ~46 GB (72% de la RAM unificada)',
            ha='center', color=C_ACCENT, fontsize=8)
    path = INFOGRAFIAS_DIR / 'fig-capstone01-01.png'
    save_fig(fig, path)
    return path


# ── Infografía 21 — Capstone 02: Pipeline 7 Agentes ─────────────────────────
def infografia_capstone02():
    fig, ax = fig_setup(7.0, 4.5)
    fig_title(ax, 'Capstone 02 — Pipeline de 7 Agentes: YouTube Shorts / TikTok')
    agents = [
        ('1\nScript\nWriter', 'LLM → guion'),
        ('2\nVoice\nSynth', 'piper-tts → WAV'),
        ('3\nImage\nGen', 'SD WebUI → frames'),
        ('4\nVideo\nAnim', 'AnimateDiff → MP4'),
        ('5\nSubtitle\nGen', 'Whisper → SRT'),
        ('6\nVideo\nEditor', 'ffmpeg → clip'),
        ('7\nUploader', 'YouTube / TikTok API'),
    ]
    bw = 1.2
    gap = 0.1
    total_w = len(agents) * bw + (len(agents) - 1) * gap
    start_x = (10 - total_w) / 2
    y_box = 4.5
    colors = [C_MID, '#0A6060', '#0A7070', '#0A5050', '#0A6A6A', '#0A4A4A', '#0A7070']
    for i, ((label, sub), color) in enumerate(zip(agents, colors)):
        x = start_x + i * (bw + gap)
        draw_box(ax, x, y_box, bw, 3.0, label, sub=sub, color=color, fontsize=7.5, subsize=6.5)
        if i < len(agents) - 1:
            draw_arrow(ax, x + bw, y_box + 1.5, x + bw + gap, y_box + 1.5)
    # N8N orquestador
    draw_box(ax, 1.5, 1.2, 7.0, 1.8, 'N8N Orquestador',
             sub='Gestión de cola · Retry · Monitoreo · Schedulers cron',
             color='#083030', fontsize=8.5)
    draw_arrow_v(ax, 5.0, 4.5, 3.0)
    path = INFOGRAFIAS_DIR / 'fig-capstone02-01.png'
    save_fig(fig, path)
    return path


# ── Mapa: chapter_id → función generadora ────────────────────────────────────
INFOGRAFIA_MAP = {
    'cap01-01': infografia_cap01_arquitectura_sistema,
    'cap02-01': infografia_cap02_force_recovery,
    'cap04-01': infografia_cap04_memoria_unificada,
    'cap05-01': infografia_cap05_rendimiento,
    'cap09-01': infografia_cap09_gpu_containers,
    'cap11-01': infografia_cap11_openclaw,
    'cap11a-01': infografia_cap11a_openclaw,
    'cap11b-01': infografia_cap11b_nemoclaw,
    'cap13-01': infografia_cap13_stt_llm_tts,
    'cap14-01': infografia_cap14_imagen_video,
    'cap15-01': infografia_cap15_pipeline_webhook,
    'cap15-02': infografia_cap15_pipeline_email,
    'cap15-03': infografia_cap15_integracion,
    'cap17-01': infografia_cap17_boot_state,
    'cap19-01': infografia_cap19_ciclo_desarrollo,
    'cap20-01': infografia_cap20_jetson_containers,
    'cap20-02': infografia_cap20_puertos,
    'cap23-01': infografia_cap23_turismo,
    'cap28-01': infografia_cap28_aiaas,
    'capstone01-01': infografia_capstone01,
    'capstone02-01': infografia_capstone02,
}


def generate_all_infografias():
    """Genera todas las infografías y retorna diccionario path→PNG."""
    INFOGRAFIAS_DIR.mkdir(exist_ok=True)
    generated = {}
    for key, fn in INFOGRAFIA_MAP.items():
        print(f'  Generando infografía: {key}')
        path = fn()
        generated[key] = path
    print(f'  {len(generated)} infografías generadas en {INFOGRAFIAS_DIR}/')
    return generated


def _get_infografia_key_from_comment(comment_text, md_path_stem):
    """Mapea el texto del comentario <!-- INFOGRAFÍA: ... --> al key del mapa."""
    # Extraer chapter id del stem del archivo
    stem = md_path_stem  # ej: 'capitulo-15-n8n'
    # Construir chapter prefix
    prefix_map = {
        'capitulo-01': 'cap01', 'capitulo-02': 'cap02',
        'capitulo-03': 'cap03', 'capitulo-04': 'cap04',
        'capitulo-05': 'cap05', 'capitulo-06': 'cap06',
        'capitulo-07': 'cap07', 'capitulo-08': 'cap08',
        'capitulo-09': 'cap09', 'capitulo-10': 'cap10',
        'capitulo-11-': 'cap11-', 'capitulo-11a': 'cap11a',
        'capitulo-11b': 'cap11b', 'capitulo-11c': 'cap11c',
        'capitulo-11d': 'cap11d', 'capitulo-12': 'cap12',
        'capitulo-13': 'cap13', 'capitulo-14': 'cap14',
        'capitulo-15': 'cap15', 'capitulo-16': 'cap16',
        'capitulo-17': 'cap17', 'capitulo-18': 'cap18',
        'capitulo-19': 'cap19', 'capitulo-20': 'cap20',
        'capitulo-21': 'cap21', 'capitulo-22': 'cap22',
        'capitulo-23': 'cap23', 'capitulo-24': 'cap24',
        'capitulo-25': 'cap25', 'capitulo-26': 'cap26',
        'capitulo-27': 'cap27', 'capitulo-28': 'cap28',
        'capstone-01': 'capstone01', 'capstone-02': 'capstone02',
    }
    chapter_prefix = 'unknown'
    for k, v in prefix_map.items():
        if k in stem:
            chapter_prefix = v
            break
    # Contar cuántas infografías hemos visto para este capítulo (se asigna en el parser)
    return chapter_prefix


# ═══════════════════════════════════════════════════════════════════════════════
# ASCII DIAGRAM DETECTION (obs#2)
# ═══════════════════════════════════════════════════════════════════════════════

_BOX_CHARS = frozenset('┌┐└┘─│├┤┬┴┼╔╗╚╝═║╠╣╦╩╬↓↑←→▼▲⟶⟵')
_TREE_RE   = re.compile(r'^[├└│]\s*[─]')


def _line_has_box(ln):
    return any(c in _BOX_CHARS for c in ln)


def add_ascii_diagram(doc, ascii_text):
    """Renderiza un diagrama ASCII como imagen matplotlib con paleta NVIDIA."""
    lines = [l for l in ascii_text.split('\n') if l.strip()]
    if not lines:
        return
    n_lines = len(lines)
    fig_h = max(2.0, min(n_lines * 0.38 + 1.2, 6.0))

    fig, ax = plt.subplots(figsize=(7.0, fig_h))
    fig.patch.set_facecolor(C_DARK)
    ax.set_facecolor(C_DARK)
    ax.axis('off')

    text_content = '\n'.join(lines)
    ax.text(0.04, 0.96, text_content,
            transform=ax.transAxes,
            ha='left', va='top',
            fontsize=8.5, fontfamily='monospace',
            color=C_WHITE,
            linespacing=1.4)

    INFOGRAFIAS_DIR.mkdir(exist_ok=True)
    key = abs(hash(ascii_text[:80])) % 100000
    path = INFOGRAFIAS_DIR / f'ascii_{key:05d}.png'
    fig.savefig(str(path), dpi=150, bbox_inches='tight', facecolor=C_DARK)
    plt.close(fig)

    try:
        doc.add_picture(str(path), width=Inches(5.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f'    [AVISO] No se pudo insertar diagrama ASCII: {e}')


# ═══════════════════════════════════════════════════════════════════════════════
# PARSER MARKDOWN EXTENDIDO
# ═══════════════════════════════════════════════════════════════════════════════

def parse_markdown(md_text, md_path_stem=''):
    """Convierte markdown a lista de bloques tipados, incluyendo infografías."""
    blocks = []
    lines = md_text.split('\n')
    i = 0
    chapter_prefix = _get_infografia_key_from_comment('', md_path_stem)
    infografia_counters = {}  # chapter_prefix → count

    while i < len(lines):
        line = lines[i]

        # Comentario HTML → infografía
        m_info = re.match(r'\s*<!--\s*INFOGRAFÍA:\s*(.+?)\s*-->', line, re.IGNORECASE)
        if m_info:
            desc = m_info.group(1).strip()
            count = infografia_counters.get(chapter_prefix, 0) + 1
            infografia_counters[chapter_prefix] = count
            key = f'{chapter_prefix}-{count:02d}'
            blocks.append({'type': 'infographic', 'key': key, 'description': desc})
            i += 1
            continue

        # Ignorar otros comentarios HTML
        if line.strip().startswith('<!--'):
            i += 1
            continue

        # Bloque de código
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip() or 'bash'
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            while code_lines and not code_lines[-1].strip():
                code_lines.pop()
            blocks.append({'type': 'code', 'lang': lang, 'content': '\n'.join(code_lines)})
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.+)', line)
        if m:
            level = len(m.group(1))
            blocks.append({'type': 'heading', 'level': level, 'text': m.group(2).strip()})
            i += 1
            continue

        # Blockquote → callout
        if line.startswith('>'):
            bq_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                bq_lines.append(lines[i][1:].strip())
                i += 1
            text = ' '.join(bq_lines)
            kind = 'NOTA'
            for kw in ('IMPORTANTE', 'ADVERTENCIA', 'ATENCIÓN', 'ATENCION', 'CONSEJO'):
                if kw in text.upper():
                    kind = kw
                    break
            text = re.sub(r'\*\*(.+?)\*\*', lambda m2: m2.group(1), text)
            text = re.sub(r'^(IMPORTANTE|NOTA|CONSEJO|ADVERTENCIA|ATENCIÓN|ATENCION):\s*', '', text, flags=re.I)
            blocks.append({'type': 'callout', 'kind': kind, 'content': text})
            continue

        # Tabla
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                def parse_row(r):
                    return [c.strip() for c in r.strip().strip('|').split('|')]
                headers = parse_row(table_lines[0])
                data_start = 2 if (len(table_lines) > 1 and re.match(r'[\s\|:\-]+', table_lines[1])) else 1
                rows = [parse_row(r) for r in table_lines[data_start:] if r.strip()]
                blocks.append({'type': 'table', 'headers': headers, 'rows': rows})
            continue

        # Lista con viñetas
        if re.match(r'^[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i]):
                items.append(re.sub(r'^[-*]\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'bullets', 'items': items, 'numbered': False})
            continue

        # Lista numerada
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i]):
                items.append(re.sub(r'^\d+\.\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'bullets', 'items': items, 'numbered': True})
            continue

        # Regla horizontal
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Diagrama ASCII / box-drawing art (obs#2) — fuera de code blocks
        if _line_has_box(line) or _TREE_RE.match(line):
            diagram_lines = [line]
            i += 1
            while i < len(lines):
                nxt = lines[i]
                if not nxt.strip():
                    # Línea vacía: continuar solo si la siguiente tiene box chars
                    if i + 1 < len(lines) and (_line_has_box(lines[i + 1]) or _TREE_RE.match(lines[i + 1])):
                        diagram_lines.append(nxt)
                        i += 1
                        continue
                    else:
                        break
                if _line_has_box(nxt) or _TREE_RE.match(nxt):
                    diagram_lines.append(nxt)
                    i += 1
                else:
                    break
            content = '\n'.join(diagram_lines).rstrip()
            if content.strip():
                blocks.append({'type': 'ascii_diagram', 'content': content})
            continue

        # Párrafo de cuerpo
        if line.strip():
            para_lines = []
            while i < len(lines) and lines[i].strip() and \
                  not lines[i].startswith('#') and \
                  not lines[i].startswith('>') and \
                  not lines[i].startswith('```') and \
                  not lines[i].strip().startswith('|') and \
                  not re.match(r'^[-*]\s+', lines[i]) and \
                  not re.match(r'^\d+\.\s+', lines[i]) and \
                  not re.match(r'^---+\s*$', lines[i]) and \
                  not lines[i].strip().startswith('<!--'):
                para_lines.append(lines[i])
                i += 1
            text = ' '.join(para_lines)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            if text.strip():
                blocks.append({'type': 'body', 'content': text.strip()})
            continue

        i += 1

    return blocks


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def _is_front_matter(md_path):
    """Prólogo e introducción usan estilo Front Matter."""
    stem = Path(md_path).stem
    return any(x in stem for x in ('prologo', 'introduccion'))


def _is_part_divider(md_path):
    """Detecta dividers de parte (00-parte-a, 00-parte-b, 00-parte-c)."""
    return '00-parte' in str(md_path)


def build_chapter_content(doc, md_path, infografias_generated=None):
    """Añade el contenido de un archivo MD al documento doc."""
    md_text = Path(md_path).read_text(encoding='utf-8')
    stem = Path(md_path).stem
    is_fm = _is_front_matter(md_path)
    is_part = _is_part_divider(md_path)

    blocks = parse_markdown(md_text, md_path_stem=stem)

    after_heading = False
    first_para_done = False

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = block['level']
            text = block['text']
            if level == 1:
                if is_part:
                    add_part_title(doc, text)
                else:
                    add_chapter_title(doc, text)
            elif level == 2:
                if is_part:
                    # subtítulo de parte
                    ps = doc.add_paragraph(style=get_style(doc, 'Heading 2'))
                    ps.paragraph_format.space_before = Pt(4)
                    ps.paragraph_format.space_after = Pt(14)
                    ps.add_run(text)
                else:
                    add_subhead(doc, text, level=2)
            elif level == 3:
                add_subhead(doc, text, level=3)
            elif level == 4:
                add_subsubhead(doc, text)
            after_heading = True
            first_para_done = False

        elif btype == 'body':
            if after_heading and not first_para_done:
                add_first_paragraph(doc, block['content'], is_front_matter=is_fm)
                first_para_done = True
            else:
                add_body_paragraph(doc, block['content'], is_front_matter=is_fm)
            after_heading = False

        elif btype == 'code':
            add_code_block(doc, block['content'], block.get('lang', 'bash'))
            after_heading = False
            first_para_done = True  # tras código, el siguiente párrafo no es "first"

        elif btype == 'callout':
            add_callout(doc, block['content'], block.get('kind', 'NOTA'))
            after_heading = False

        elif btype == 'table':
            if block.get('headers') and block.get('rows'):
                add_table(doc, block['headers'], block['rows'])
            after_heading = False

        elif btype == 'bullets':
            add_bullets(doc, block['items'], block.get('numbered', False), is_front_matter=is_fm)
            after_heading = False

        elif btype == 'ascii_diagram':
            add_ascii_diagram(doc, block['content'])  # obs#2
            after_heading = False

        elif btype == 'infographic':
            key = block['key']
            if infografias_generated and key in infografias_generated:
                png_path = infografias_generated[key]
                if png_path.exists():
                    try:
                        doc.add_picture(str(png_path), width=Inches(5.5))
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Caption corto a partir de la descripción
                        desc = block['description']
                        short_desc = desc.split('—')[0].strip()
                        if len(short_desc) > 80:
                            short_desc = short_desc[:77] + '...'
                        add_image_caption(doc, short_desc)
                    except Exception as e:
                        print(f'    [AVISO] No se pudo insertar imagen {key}: {e}')
                        add_callout(doc, f'[Figura: {short_desc}]', 'NOTA')
            else:
                # Placeholder si la infografía no se generó
                add_callout(doc, f'[INFOGRAFÍA PENDIENTE — {block["description"][:100]}]', 'NOTA')
            after_heading = False


def build_chapter_docx(md_path, output_path, infografias_generated=None):
    """Genera un DOCX individual para un capítulo."""
    doc = create_chapter_doc()
    build_chapter_content(doc, md_path, infografias_generated)
    doc.save(str(output_path))
    print(f'  OK: {Path(output_path).name}')


def build_all_chapters(infografias_generated=None):
    """Genera 42 DOCX individuales en book/docx-por-capitulo/."""
    DOCX_OUT_DIR.mkdir(exist_ok=True)
    count = 0
    for rel_path in BOOK_CHAPTERS:
        md_path = BOOK_DIR / rel_path
        if not md_path.exists():
            print(f'  [AVISO] No encontrado: {rel_path}')
            continue
        chapter_name = md_path.stem
        out_path = DOCX_OUT_DIR / f'{chapter_name}.docx'
        build_chapter_docx(md_path, out_path, infografias_generated)
        count += 1
    print(f'\n{count} capítulos generados en {DOCX_OUT_DIR}/')
    return count


def verify_chapter_refs():
    """obs#8: Verifica que todas las referencias 'Capítulo X' en los .md coincidan con capítulos reales."""
    import re as _re
    # Construir mapa número → título desde el #-heading de cada capítulo
    cap_map = {}
    cap_re = _re.compile(r'[Cc]ap(?:ítulo|\.)\s*(\d+[A-Za-z]?)', _re.IGNORECASE)
    for rel in BOOK_CHAPTERS:
        md_path = BOOK_DIR / rel
        if not md_path.exists():
            continue
        try:
            first_line = next(
                (l.lstrip('#').strip() for l in md_path.read_text(encoding='utf-8').splitlines()
                 if l.startswith('#') and not l.startswith('##')), '')
            m = _re.match(r'[Cc]ap(?:ítulo|\.)\s*(\d+[A-Za-z]?)', first_line, _re.IGNORECASE)
            if m:
                cap_map[m.group(1).upper()] = first_line
        except Exception:
            pass

    print(f'Capítulos indexados: {sorted(cap_map.keys())}')
    issues = []
    for rel in BOOK_CHAPTERS:
        md_path = BOOK_DIR / rel
        if not md_path.exists():
            continue
        try:
            text = md_path.read_text(encoding='utf-8')
            in_code = False
            for ln, line in enumerate(text.splitlines(), 1):
                if line.strip().startswith('```'):
                    in_code = not in_code
                if in_code:
                    continue
                for m in cap_re.finditer(line):
                    num = m.group(1).upper()
                    if num not in cap_map:
                        issues.append(f'  {md_path.name}:{ln} — "{m.group(0)}" (num {num} no encontrado)')
        except Exception:
            pass

    if issues:
        print(f'\n{len(issues)} referencias potencialmente incorrectas:')
        for issue in issues:
            print(issue)
    else:
        print('\nTodas las referencias de capítulo son válidas.')
    return issues


def compile_manuscript(output_path, infografias_generated=None):
    """Compila todos los capítulos en un único DOCX manuscrito.
    obs#10: usa add_page_break() — NO add_section() — para mantener una única sección
    compartiendo encabezados, pies y numeración de página en todo el manuscrito."""
    doc = create_chapter_doc()
    first = True
    for rel_path in BOOK_CHAPTERS:
        md_path = BOOK_DIR / rel_path
        if not md_path.exists():
            print(f'  [AVISO] No encontrado: {rel_path}')
            continue
        if not first:
            doc.add_page_break()   # obs#10: page break, NOT section break
        first = False
        print(f'  Compilando: {rel_path}')
        build_chapter_content(doc, md_path, infografias_generated)
    doc.save(str(output_path))
    print(f'\nManuscrito compilado: {output_path}')


# ═══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(f'Uso: python {sys.argv[0]} <archivo.md>')
        print(f'     python {sys.argv[0]} --batch')
        print(f'     python {sys.argv[0]} --compile [salida.docx]')
        print(f'     python {sys.argv[0]} --verify-refs')
        sys.exit(1)

    mode = sys.argv[1]

    if mode == '--verify-refs':
        issues = verify_chapter_refs()
        sys.exit(1 if issues else 0)

    if mode == '--batch':
        print('Generando infografías...')
        infos = generate_all_infografias()
        print(f'\nGenerando {len(BOOK_CHAPTERS)} DOCX individuales...')
        build_all_chapters(infos)
        sys.exit(0)

    if mode == '--compile':
        out_file = sys.argv[2] if len(sys.argv) >= 3 else str(BOOK_DIR / 'Jetson_AGX_Orin_JP72_KDP_FINAL.docx')
        print('Generando infografías...')
        infos = generate_all_infografias()
        print(f'\nCompilando manuscrito → {out_file}')
        compile_manuscript(out_file, infos)
        sys.exit(0)

    # Capítulo individual
    md_file = Path(mode)
    if not md_file.exists():
        print(f'Error: no se encuentra {md_file}')
        sys.exit(1)
    out_file = sys.argv[2] if len(sys.argv) >= 3 else str(md_file.with_suffix('.docx'))
    print('Generando infografías...')
    infos = generate_all_infografias()
    print(f'\nGenerando DOCX: {md_file.name}')
    build_chapter_docx(md_file, out_file, infos)
