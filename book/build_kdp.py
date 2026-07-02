#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_kdp.py — Convierte capítulos Markdown a DOCX con formato Amazon KDP (6×9 pulgadas).

Uso (capítulo individual):
    python book/build_kdp.py book/parte-a/capitulo-10-motores-inferencia/capitulo-10-motores-inferencia.md

Uso (libro completo en un solo DOCX):
    python book/build_kdp.py --all [salida.docx]
    # Por defecto genera: book/Jetson_AGX_Orin_JP72_Complete_Guide.docx

Estructura del libro:
    book/parte-a/   → Capítulos 1–17 (Fundamentos + Stack IA)
    book/parte-b/   → Capítulos 13img, 18–26 (Proyectos Prácticos)
    book/parte-c/   → Capítulo 27, Capstones 01–02, Conclusiones
    book/glosario/  → Glosario

Requiere: pip install python-docx
"""

import sys
import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta KDP (optimizada para impresión B&W y pantalla) ─────────────────────
ACCENT       = '1A5A6A'   # Encabezados y acentos (azul oscuro — visible en B&W)
TABLE_HDR    = 'CCCCCC'   # Fondo encabezado tabla (gris claro — B&W safe)
CODE_BG      = 'F2F2F2'   # Fondo código (gris muy claro — imprime bien)
CODE_FG      = '1A1A1A'   # Texto código (casi negro)
ALT_ROW      = 'F8F8F8'   # Fila alternante (gris muy tenue)
CALLOUT_BDR  = 'AAAAAA'   # Borde cajas de aviso

# Tamaño de página KDP 6"×9"
PAGE_W = Inches(6.0)
PAGE_H = Inches(9.0)

# Márgenes KDP (para libro de 300-500 páginas)
MARGIN_TOP    = Inches(0.75)
MARGIN_BOTTOM = Inches(0.75)
MARGIN_INNER  = Inches(0.875)  # gutter (encuadernación)
MARGIN_OUTER  = Inches(0.625)


def hex_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.lstrip('#'))
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), kwargs[side].get('val', 'single'))
            el.set(qn('w:sz'), str(kwargs[side].get('sz', 4)))
            el.set(qn('w:color'), kwargs[side].get('color', '000000'))
            tc_borders.append(el)
    tc_pr.append(tc_borders)


def table_full_width(table, doc):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def create_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = PAGE_W
    sec.page_height   = PAGE_H
    sec.left_margin   = MARGIN_INNER
    sec.right_margin  = MARGIN_OUTER
    sec.top_margin    = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM

    # Estilos base
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(13.2)  # 1.15 × 11pt

    return doc


def add_heading(doc, text, level):
    """level 1=parte, 2=sección, 3=subsección"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 8)
    p.paragraph_format.space_after  = Pt(6 if level > 1 else 10)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    sizes = {1: 18, 2: 14, 3: 12}
    r.font.size = Pt(sizes.get(level, 11))
    if level <= 2:
        r.font.color.rgb = hex_rgb(ACCENT)
    return p


def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph(text.strip())
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code_block(doc, code, lang='bash'):
    # Outer table for background + left border
    tbl = doc.add_table(rows=1, cols=1)
    table_full_width(tbl, doc)
    cell = tbl.cell(0, 0)
    cell_bg(cell, CODE_BG)

    # Left border accent
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'thick')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), ACCENT)
    borders.append(left)
    tc_pr.append(borders)

    # Cell padding
    tc_mar = OxmlElement('w:tcMar')
    for side, val in [('top', 80), ('left', 160), ('bottom', 80), ('right', 120)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    first = True
    for line in code.split('\n'):
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Detect comment lines
        stripped = line.strip()
        is_comment = stripped.startswith('#') and lang in ('bash', 'sh', 'shell', '')
        is_comment = is_comment or (stripped.startswith('//') and lang in ('js', 'javascript', 'json', 'c', 'cpp'))

        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = hex_rgb('888888' if is_comment else CODE_FG)

    doc.add_paragraph()  # spacing after


def add_callout(doc, text, kind='NOTA'):
    kind = kind.upper().strip()
    labels = {
        'IMPORTANTE': 'IMPORTANTE',
        'NOTA':       'NOTA',
        'CONSEJO':    'CONSEJO',
        'ADVERTENCIA': 'ADVERTENCIA',
        'ATENCION':   'ATENCIÓN',
        'ATENCIÓN':   'ATENCIÓN',
    }
    label = labels.get(kind, kind)

    tbl = doc.add_table(rows=1, cols=1)
    table_full_width(tbl, doc)
    cell = tbl.cell(0, 0)
    cell_bg(cell, 'F5F5F5')

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single' if side == 'left' else 'single')
        el.set(qn('w:sz'), '12' if side == 'left' else '4')
        el.set(qn('w:color'), '444444')
        borders.append(el)
    tc_pr.append(borders)

    tc_mar = OxmlElement('w:tcMar')
    for s, v in [('top', 80), ('left', 160), ('bottom', 80), ('right', 120)]:
        m = OxmlElement(f'w:{s}')
        m.set(qn('w:w'), str(v))
        m.set(qn('w:type'), 'dxa')
        tc_mar.append(m)
    tc_pr.append(tc_mar)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r_label = p.add_run(f'{label}: ')
    r_label.bold = True
    r_label.font.name = 'Times New Roman'
    r_label.font.size = Pt(10)
    r_body = p.add_run(text.strip())
    r_body.font.name = 'Times New Roman'
    r_body.font.size = Pt(10)

    doc.add_paragraph()


def add_table(doc, headers, rows):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    table_full_width(tbl, doc)

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell_bg(cell, TABLE_HDR)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        tbl_row = tbl.rows[ri + 1]
        bg = ALT_ROW if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            if ci >= n_cols:
                break  # skip extra columns that exceed header count
            cell = tbl_row.cells[ci]
            cell_bg(cell, bg)
            p = cell.paragraphs[0]
            # Strip markdown bold/italic
            val_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', str(val))
            val_clean = re.sub(r'\*(.+?)\*', r'\1', val_clean)
            r = p.add_run(val_clean)
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'

    doc.add_paragraph()


def add_bullets(doc, items, numbered=False):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        prefix = f'{i}. ' if numbered else '• '
        r = p.add_run(prefix + item.strip())
        r.font.size = Pt(10.5)
        r.font.name = 'Times New Roman'
    doc.add_paragraph()


# ── Parser de Markdown ────────────────────────────────────────────────────────

def parse_markdown(md_text):
    """Convierte markdown a una lista de bloques tipados."""
    blocks = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip() or 'bash'
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # Remove trailing empty lines
            while code_lines and not code_lines[-1].strip():
                code_lines.pop()
            blocks.append({'type': 'code', 'lang': lang, 'content': '\n'.join(code_lines)})
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.+)', line)
        if m:
            level = len(m.group(1))
            blocks.append({'type': 'heading', 'level': level, 'text': m.group(2).strip()})
            i += 1
            continue

        # Blockquote → callout
        if line.startswith('>'):
            text = line[1:].strip()
            kind = 'NOTA'
            for kw in ('IMPORTANTE', 'ADVERTENCIA', 'ATENCIÓN', 'ATENCION', 'CONSEJO'):
                if kw in text.upper():
                    kind = kw
                    break
            # Remove bold markers and known callout prefixes
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'^(IMPORTANTE|NOTA|CONSEJO|ADVERTENCIA|ATENCIÓN|ATENCION):\s*', '', text, flags=re.I)
            blocks.append({'type': 'callout', 'kind': kind, 'content': text})
            i += 1
            continue

        # Table
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Parse table
            if len(table_lines) >= 2:
                def parse_row(r):
                    return [c.strip() for c in r.strip().strip('|').split('|')]
                headers = parse_row(table_lines[0])
                # Skip separator row (---)
                data_start = 1
                if len(table_lines) > 1 and re.match(r'[\s\|:\-]+', table_lines[1]):
                    data_start = 2
                rows = [parse_row(r) for r in table_lines[data_start:]]
                blocks.append({'type': 'table', 'headers': headers, 'rows': rows})
            continue

        # Bullet list
        if re.match(r'^[-*]\s+', line):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i]):
                items.append(re.sub(r'^[-*]\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'bullets', 'items': items, 'numbered': False})
            continue

        # Numbered list
        if re.match(r'^\d+\.\s+', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i]):
                items.append(re.sub(r'^\d+\.\s+', '', lines[i]))
                i += 1
            blocks.append({'type': 'bullets', 'items': items, 'numbered': True})
            continue

        # Horizontal rule
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Body paragraph (accumulate non-empty lines)
        if line.strip():
            para_lines = []
            while i < len(lines) and lines[i].strip() and \
                  not lines[i].startswith('#') and \
                  not lines[i].startswith('>') and \
                  not lines[i].startswith('```') and \
                  not (lines[i].strip().startswith('|')) and \
                  not re.match(r'^[-*]\s+', lines[i]) and \
                  not re.match(r'^\d+\.\s+', lines[i]) and \
                  not re.match(r'^---+\s*$', lines[i]):
                para_lines.append(lines[i])
                i += 1
            text = ' '.join(para_lines)
            # Clean markdown inline formatting
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            if text.strip():
                blocks.append({'type': 'body', 'content': text.strip()})
            continue

        i += 1

    return blocks


def build_docx(md_path, output_path):
    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_markdown(md_text)

    doc = create_doc()

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = block['level']
            text  = block['text']
            if level == 1:
                add_heading(doc, text, 1)
            elif level == 2:
                add_heading(doc, text, 2)
            else:
                add_heading(doc, text, 3)

        elif btype == 'body':
            add_body(doc, block['content'])

        elif btype == 'code':
            add_code_block(doc, block['content'], block.get('lang', 'bash'))

        elif btype == 'callout':
            add_callout(doc, block['content'], block.get('kind', 'NOTA'))

        elif btype == 'table':
            if block.get('headers') and block.get('rows'):
                add_table(doc, block['headers'], block['rows'])

        elif btype == 'bullets':
            add_bullets(doc, block['items'], block.get('numbered', False))

    doc.save(output_path)
    print(f'DOCX generado: {output_path}')


# Orden canónico del libro completo (relativo a book/)
BOOK_CHAPTERS = [
    'parte-a/capitulo-01-introduccion/capitulo-01-introduccion.md',
    'parte-a/capitulo-02-primer-arranque/capitulo-02-primer-arranque.md',
    'parte-a/capitulo-03-configuracion-base/capitulo-03-configuracion-base.md',
    'parte-a/capitulo-04-memoria-almacenamiento/capitulo-04-memoria-almacenamiento.md',
    'parte-a/capitulo-05-rendimiento/capitulo-05-rendimiento.md',
    'parte-a/capitulo-06-shell-entorno/capitulo-06-shell-entorno.md',
    'parte-a/capitulo-07-red/capitulo-07-red.md',
    'parte-a/capitulo-08-acceso-remoto/capitulo-08-acceso-remoto.md',
    'parte-a/capitulo-09-docker/capitulo-09-docker.md',
    'parte-a/capitulo-10-motores-inferencia/capitulo-10-motores-inferencia.md',
    'parte-a/capitulo-11-stack-agentico/capitulo-11-stack-agentico.md',
    'parte-a/capitulo-11a-openclaw/capitulo-11a-openclaw.md',
    'parte-a/capitulo-11b-nemoclaw/capitulo-11b-nemoclaw.md',
    'parte-a/capitulo-11c-open-webui/capitulo-11c-open-webui.md',
    'parte-a/capitulo-11d-tool-calling/capitulo-11d-tool-calling.md',
    'parte-a/capitulo-12-computer-vision/capitulo-12-computer-vision.md',
    'parte-b/capitulo-13-imagen-video/capitulo-13-imagen-video.md',
    'parte-a/capitulo-14-n8n/capitulo-14-n8n.md',
    'parte-a/capitulo-15-benchmarking/capitulo-15-benchmarking.md',
    'parte-a/capitulo-16-produccion/capitulo-16-produccion.md',
    'parte-a/capitulo-17-troubleshooting/capitulo-17-troubleshooting.md',
    'parte-b/capitulo-18-python-vscode/capitulo-18-python-vscode.md',
    'parte-b/capitulo-19-jetson-containers/capitulo-19-jetson-containers.md',
    'parte-b/capitulo-20-pdf-podcast/capitulo-20-pdf-podcast.md',
    'parte-b/capitulo-21-transcripcion-reuniones/capitulo-21-transcripcion-reuniones.md',
    'parte-b/capitulo-22-agencia-turismo/capitulo-22-agencia-turismo.md',
    'parte-b/capitulo-23-embudo-ventas/capitulo-23-embudo-ventas.md',
    'parte-b/capitulo-24-linkedin/capitulo-24-linkedin.md',
    'parte-b/capitulo-25-asistente-voz/capitulo-25-asistente-voz.md',
    'parte-b/capitulo-26-rag/capitulo-26-rag.md',
    'parte-c/capitulo-27-microservicios/capitulo-27-microservicios.md',
    'parte-c/capstone-01-agencia-ia/capstone-01-agencia-ia.md',
    'parte-c/capstone-02-automatizacion-video/capstone-02-automatizacion-video.md',
    'parte-c/conclusiones/conclusiones.md',
    'glosario/glosario.md',
    'parte-a/capitulo-xx-tts-stt/capitulo-xx-tts-stt.md',  # ubicación pendiente
    'parte-a/appendix/appendix.md',
]


def build_all_docx(book_dir, output_path):
    """Genera un DOCX con todos los capítulos del libro en orden canónico."""
    book_root = Path(book_dir)
    doc = create_doc()

    for rel_path in BOOK_CHAPTERS:
        md_path = book_root / rel_path
        if not md_path.exists():
            print(f'  [AVISO] No encontrado: {rel_path}')
            continue
        print(f'  Procesando: {rel_path}')
        md_text = md_path.read_text(encoding='utf-8')
        blocks = parse_markdown(md_text)
        for block in blocks:
            btype = block['type']
            if btype == 'heading':
                level = min(block['level'], 3)
                add_heading(doc, block['text'], level)
            elif btype == 'body':
                add_body(doc, block['content'])
            elif btype == 'code':
                add_code_block(doc, block['content'], block.get('lang', 'bash'))
            elif btype == 'callout':
                add_callout(doc, block['content'], block.get('kind', 'NOTA'))
            elif btype == 'table':
                if block.get('headers') and block.get('rows'):
                    add_table(doc, block['headers'], block['rows'])
            elif btype == 'bullets':
                add_bullets(doc, block['items'], block.get('numbered', False))
        # Page break between chapters
        doc.add_page_break()

    doc.save(output_path)
    print(f'\nLibro completo generado: {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(f'Uso: python {sys.argv[0]} <archivo.md> [salida.docx]')
        print(f'     python {sys.argv[0]} --all [salida.docx]')
        sys.exit(1)

    if sys.argv[1] == '--all':
        book_dir = Path(__file__).parent
        out_file = sys.argv[2] if len(sys.argv) >= 3 else str(book_dir / 'Jetson_AGX_Orin_JP72_Complete_Guide.docx')
        build_all_docx(book_dir, out_file)
        sys.exit(0)

    md_file = sys.argv[1]
    if len(sys.argv) >= 3:
        out_file = sys.argv[2]
    else:
        out_file = str(Path(md_file).with_suffix('.docx'))

    if not Path(md_file).exists():
        print(f'Error: no se encuentra {md_file}')
        sys.exit(1)

    build_docx(md_file, out_file)
